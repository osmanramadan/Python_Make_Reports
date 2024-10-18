# Import Required Libraries
## PyQt6 Imports
from PyQt6.QtWidgets import (
    QApplication, QGridLayout, QHBoxLayout, QVBoxLayout, 
    QSplitter, QWidget, QPushButton, QLineEdit, QLabel, 
    QMessageBox, QItemDelegate, QTextEdit, QFrame, 
    QFileDialog, QScrollArea, QMainWindow, 
    QTableWidget, QTableWidgetItem, QRadioButton, 
    QProgressBar, QListWidget, QListWidgetItem
)
from PyQt6.QtGui import QIcon, QFont, QPixmap
from PyQt6 import uic
from PyQt6.QtCore import Qt, pyqtSignal
## Standard Library Imports
import sqlite3
import sys
import os
import webbrowser
import shutil
import ctypes
import time
## Third-Party Imports
from PIL import Image, ImageOps
import pyautogui
import docx
import convert_numbers
import arabic_reshaper
from bidi.algorithm import get_display
## Docx Imports
import docx.enum
import docx.enum.section
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.shared import qn
from docx.oxml.ns import qn as qn2
## ReportLab Imports
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape, letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table,TableStyle, Paragraph, Image as img
from reportlab.lib.enums import TA_LEFT, TA_RIGHT
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle




xForImpo = 900
yForImpo = 0
title = "موثق البرامج"
desktopPath = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')


def is_admin():
    try:
        return ctypes.windll.shell32.IsUserAnAdmin()
    except:
        return False
    
if is_admin():       
    class LoadingPage(QWidget):
        def __init__(self):
            super().__init__()
            uic.loadUi("design/loadingPage.ui",self)
    class LineEditDelegate(QItemDelegate):
        def createEditor(self,parent,option,index):
            editor = QLineEdit(parent)
            editor.setFrame(False)
            editor.setReadOnly(True)
            editor.setAlignment(Qt.AlignmentFlag.AlignCenter)
            return editor
    class Choices(QWidget):
        def __init__(self):
            super().__init__()
            self.setWindowTitle(title)
            self.setWindowIcon(QIcon("icons/icon.ico"))
    class ClickableQFrame(QFrame):
        clicked = pyqtSignal()
        def mousePressEvent(self, event):
            self.clicked.emit()

    # Init variable to connect to database 
    con = sqlite3.connect("app.db")
    cr = con.cursor()
    class ReportEditor(QMainWindow):
        def __init__(self):
            super().__init__()
            self.setWindowTitle('Report Editor')
            self.setGeometry(100, 100, 1200, 800)
            self.setWindowIcon(QIcon('icons/icon.ico'))
            self.initUI()
            self.show()
        # Init App Variables
        def initUI(self):
            self.ablePrograme= False
            self.ableGoals= False
            self.ableDescription= False
            self.ableCreator= False
            self.ableDate= False
            self.ableBenefits= False
            self.ableCount= False
            self.countPic = 0
            self.secretLittleThing = ""
            global yForImpo
            global xForImpo
            yForImpo = 50
            xForImpo = 0
            self.pictersPaths = ["","","",""]
            self.windowCreating = Choices()
            self.windowCreating.setMinimumSize(900,1150)
            self.setMinimumSize(930,500)
            self.setWindowTitle(title)
            self.showMaximized()
            self.windowCreating.setWindowTitle(title)
            self.windowCreating.setWindowIcon(QIcon("icons/icon.ico"))
            button_list_widget = QWidget(self.windowCreating)
            button_list_layout = QVBoxLayout()            
            button_list_widget.setLayout(button_list_layout)
            button_list_widget.setMaximumWidth(200)   

            #Function To Add App Buttons
            def addButton(icon, text, callback):
                button = QPushButton(text)     
                button.setStyleSheet(
                                        f"""
                                        font-size: 14px;
                                        qproperty-icon: url('{icon}');
                                        qproperty-iconSize: 25px 25px;
                                        background-color: #2ABCB5;   
                                        cursor: pointer !important;
                                        """
                                    )
                button.setLayoutDirection(Qt.LayoutDirection.RightToLeft)
                button.setMinimumHeight(40) 
                button.clicked.connect(callback)           
                button.setCursor(Qt.CursorShape.PointingHandCursor)
                button_list_layout.addWidget(button)     
            # Add buttons to the layout instead of toolbar
            addButton("images/addNewreport.png", "إنشاء تقرير", self.createReportFun)
            addButton("images/savedReports.png", "التقارير المحفوظة", self.savedReportsFun)
            addButton("images/save.png", "حفظ باسم", self.SavePrograme)
            addButton("images/clearFields.png", "إفراغ الحقول", self.emptyFieldsFun)
            addButton("images/summary.png", "ملخص التقارير", self.summaryReports)
            addButton("images/printer.png", "الطباعة", self.printDoc)
            addButton("images/control.png", "لوحة التحكم", self.controlPanel)
            addButton("images/dbExport.png", "تصدير قاعدة البيانات" , self.exportDb)
            addButton("images/importDb.png", "إستيراد قاعدة البيانات", self.importDb)
            addButton("images/word.png", "word حفظ بصيغة", self.writeWord)
            addButton("images/pdfIcon.png", "pdf حفظ بصيغة", lambda: self.writePdf(fromWhere="convert"))

            companyLogobutton = QPushButton()                
            companyLogobutton.setStyleSheet(f"font-size:14px;qproperty-icon:url('images/companyLogo.png');qproperty-iconSize:150px 100px;background-color:transparent;")   
            companyLogobutton.clicked.connect(self.openwebSite)     
            companyLogobutton.setCursor(Qt.CursorShape.PointingHandCursor)    
            button_list_layout.addWidget(companyLogobutton) 
            report_widget = QWidget(self.windowCreating)


            # ******************* middle top section ******************* ******************* middle top section  *******************
            # Create a vertical layout to hold the header label and the list widget
            report_widgetlayout = QVBoxLayout()                        
            report_widget.setLayout(report_widgetlayout)  
            report_widget.setMinimumHeight(1150)
            report_widget.setMinimumSize(920,1150)
            self.hiderFrameshow = QFrame(report_widget)
            self.hiderFrameshow.setStyleSheet("background-color: white")

            self.hidderFramePicshow = ClickableQFrame(self.hiderFrameshow)
            self.hidderFramePicshow.setStyleSheet(f"background-color:#EBEAE9;")
            self.hidderFramePicshow.setGeometry(40,5,250,130)
            self.hidderFramePicshow.setCursor(Qt.CursorShape.PointingHandCursor)
            self.hidderlayoutPicshow = QVBoxLayout()
            self.hidderFramePicshow.setLayout(self.hidderlayoutPicshow)
                                                
            self.hidderFramePicshow.clicked.connect(lambda: self.putImage(f"ReportCover"))
            specialButton = QPushButton(self.hiderFrameshow)
            specialButton.setIcon(QIcon("images/cam.png"))
            specialButton.clicked.connect(lambda:self.putImage(f"ReportCover"))
            specialButton.setCursor(Qt.CursorShape.PointingHandCursor)
            specialButton.move(290,65)

            DeleteButtonHidder = QPushButton(self.hiderFrameshow)
            DeleteButtonHidder.setStyleSheet(f"Qproperty-icon:url(images/trashicon.png);qproperty-iconSize:15px 15px;background-color:rgb(253, 253, 253)")
            DeleteButtonHidder.clicked.connect(lambda:self.deleteImagesTemp(f"ReportCover"))
            DeleteButtonHidder.setCursor(Qt.CursorShape.PointingHandCursor)
            DeleteButtonHidder.move(290,90)


            Frame_text = QFrame(self.hiderFrameshow)
            Frame_text.setStyleSheet("background-color: white")
            Frame_text.setGeometry(655,0,250,140)
            text_layout = QVBoxLayout(Frame_text)

            Frame_text.setLayout(text_layout)
            cr.execute("SELECT line1 FROM start")
            Label1 = QLabel(cr.fetchone()[0])
            text_layout.addWidget(Label1)
            
            cr.execute("SELECT line2 FROM start")
            Label2 = QLabel("   "+cr.fetchone()[0])
            text_layout.addWidget(Label2)

            cr.execute("SELECT line3 FROM start")
            Label3 = QLabel(cr.fetchone()[0])
            text_layout.addWidget(Label3)

            cr.execute("SELECT line4 FROM start")

            
            Label4 = QLabel(cr.fetchone()[0])
            text_layout.addWidget(Label4)

            logoLabel = QLabel(self.hiderFrameshow)
            logoLabel.move(350,10)

            cr.execute("SELECT icon FROM start")
            
            try:
                with open("images/logo.png","wb") as logoImpo:
                  logoImpo.write(cr.fetchone()[0])
                img = Image.open("images/logo.png")
                img = img.resize((260,125),Image.LANCZOS)
                img.save("images/logo.png",quality=100)
                pix = QPixmap("images/logo.png")
                logoLabel.setPixmap(pix)
            except:
                pass

            self.hiderFrameshow.setGeometry(0,3,900,140)
            self.cFrameshow = QFrame(report_widget)
            self.cFrameshow.setStyleSheet("background-color: white")
            
            labelGood = QLabel("توثيق برنامج",self.cFrameshow)
            labelGood.setStyleSheet("font-size:20px")
            labelGood.move(410,20)

            self.ablePrograme = True
            self.createNamePrograme()
            self.ableGoals = True
            self.createGoals()
            self.ableDescription = True
            self.createDescription()
            self.ableCreator = True
            self.executer()
            self.ableDate = True
            self.executeDate()
            self.ableBenefits = True
            self.Benefits()
            self.ableCount = True
            self.CountBenefits()
            self.countPic = 4
            self.CreatePic(4)
            self.addAdmins(self.cFrameshow)            
            self.cFrameshow.setGeometry(0,145,900,1150)
            layout_widget = QWidget(self.windowCreating)
            layout_widget.setMaximumWidth(240)


            # ******************* left top section ************************* left top section *******************
            
            layout = QVBoxLayout()                        
            layout_widget.setLayout(layout)                        
            header_label = QLabel("التقارير المحفوظة")
            header_label.setAlignment(Qt.AlignmentFlag.AlignCenter)             
            header_label.setStyleSheet("font-size: 16px; font-weight: bold; padding: 10px; background-color: #2ABCB5; color:white")
            layout.addWidget(header_label)
            self.listWidget = QListWidget(layout_widget)
            layout.addWidget(self.listWidget)    
            self.load_data()          
            
            # Connect click event to open the report
            self.listWidget.clicked.connect(self.onItemClicked)
            self.listWidget.setCursor(Qt.CursorShape.PointingHandCursor)
            layout_widget.setMinimumWidth(380)
                                    
            
            self.scroll = QScrollArea()
            #Scroll Area Properties
            self.scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOn)
            self.scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOn)
            self.scroll.setWidgetResizable(True)
            self.scroll.setWidget(report_widget)
            Gridlayout_widget = QWidget()
            # Create a grid layout
            Gridlayout = QGridLayout()
            Gridlayout_widget.setLayout(Gridlayout)
            Gridlayout.addWidget(self.scroll,0, 0)
            Gridlayout.addWidget(button_list_widget,0, 1)            
            # Align the buttons to expand vertically (row stretch)
            Gridlayout.setRowStretch(0, 1)
            
            splitter = QSplitter(Qt.Orientation.Horizontal, self.windowCreating )            
            splitter.setHandleWidth(5)
            splitter.setStyleSheet("""
                QSplitter::handle:horizontal {
                    background-color: gray; 
                }
            """)
            splitter.addWidget(layout_widget)
            splitter.addWidget(Gridlayout_widget)
            self.setCentralWidget(splitter) 

            
        def openwebSite(self):
            webbrowser.open('https://www.ersal-m.com', new=2)

        # export current database
        def exportDb(self):

            FileNameSave = QFileDialog.getSaveFileName(self.windowCreating,"اختر مسارا",desktopPath)    
            if len(FileNameSave[0])>0:
                folder = (str(FileNameSave[0]).split(f"/"))
                nameFile = folder[-1]
                folderFinle = "/".join(folder[:-1])
                database_name= nameFile+".db"
                destination_path=f"{folderFinle}/{database_name}"
                if os.path.exists(destination_path):
                    confirm = QMessageBox.question(
                       self,
                      "تنبيه",
                      f"الملف موجود بالفعل هل تريد اعادة انشائه؟",
                      QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                      )
                    if confirm == QMessageBox.StandardButton.No:
                        return 

            try:
              # Copy the database file to the specified location with the new name
              shutil.copy2("app.db", destination_path) 
              # Create a custom message box with added space
              d = QMessageBox(parent=self.windowCreating,text=f"تم تصدير قاعدة البيانات '{database_name}' بنجاح")
              d.setWindowTitle("نجاح")
              d.setIcon(QMessageBox.Icon.Information)
              d.exec()
             
            except Exception as e:
              QMessageBox.critical(self.windowCreating, "فشل الاضافه: {str(e)}")

        # Import database (note:DB should be created first by developer)
        def importDb(self):

            d = QMessageBox(parent=self.windowCreating,text="هل انت متأكد من استرداد قاعدة بيانات؟")
            d.setIcon(QMessageBox.Icon.Information)
            d.setWindowTitle(title)
            d.setStandardButtons(QMessageBox.StandardButton.Cancel|QMessageBox.StandardButton.Ok)
            important = d.exec()
            if important == QMessageBox.StandardButton.Ok:
                self.completeImportDb()
        
        def completeImportDb(self):

            fileDbUser = QFileDialog.getOpenFileName(self.windowCreating,"اختر ملفا",desktopPath,filter="Database File (*.db)")
            if len(fileDbUser[0]) > 0:
                try:
                    self.con1 = sqlite3.connect(fileDbUser[0])
                    cr1 = self.con1.cursor()
                    cr1.execute("SELECT useAble FROM confirmationDatabase")
                    if cr1.fetchone()[0] == "canUse":
                        cr1.execute("SELECT * FROM reports")
                        for i in cr1.fetchall():
                            cr.execute("""INSERT INTO reports (reportName , name , Goals , description , executer , executeDate , benefits , countBenefits , pic1 , pic2 , pic3 , pic4 , picLogo , label1Maybe , label2Maybe , manger , co_manger) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""" , (i[1],i[2],i[3],i[4],i[5],i[6],i[7],i[8],i[9],i[10],i[11],i[12],i[13],i[14],i[15],i[16],i[17]))                          
                        con.commit()
                        con.close()
                        self.con1.close()
                        d = QMessageBox(parent=self.windowCreating,text="تم استيراد قاعدة البيانات بنجاح")
                        d.setWindowTitle("نجاح")
                        d.setIcon(QMessageBox.Icon.Information)
                        d.exec()
                        os.execv(sys.executable, ['python'] + sys.argv)
                        app.closeAllWindows()
                    else:
                        raise Exception("notUseAble")
                except Exception as e:
                    self.con1.close()
                    d = QMessageBox(parent=self.windowCreating,text="قاعدة البيانات غير صالحة")
                    d.setWindowTitle("ERROR")
                    d.setIcon(QMessageBox.Icon.Critical)


        def savedReportsFun(self):

            self.windowSaved = Choices()
            self.windowSaved.setFixedSize(600, 500)
            self.windowSaved.setWindowTitle(title)
            self.windowSaved.setWindowIcon(QIcon("icons/icon.ico"))
            self.savedReports = QTableWidget(self.windowSaved)
            self.savedReports.setGeometry(15, 5, 570, 450)
            self.savedReports.setColumnCount(4)
            self.savedReports.setColumnHidden(0, True)
            self.savedReports.setColumnWidth(0, 70)
            self.savedReports.setColumnWidth(1, 70)
            self.savedReports.setColumnWidth(2, 70)
            self.savedReports.setColumnWidth(3, 410)
            self.savedReports.setHorizontalHeaderLabels(["", "", "", "اسم التقرير"])

            cr.execute("SELECT id, reportName FROM reports")
            for n, i in enumerate(cr.fetchall()):
               self.savedReports.insertRow(self.savedReports.rowCount())
               # Create delete button
               button = QPushButton()
               button.setStyleSheet("QProperty-icon:url(images/trashicon.png); qproperty-iconSize:30px 30px; background-color:rgb(253, 253, 253)")
               button.clicked.connect(lambda _, row=n: self.deleteReport(row))
               button.setCursor(Qt.CursorShape.PointingHandCursor)
               self.savedReports.setIndexWidget(self.savedReports.model().index(n, 1), button)
               # Set report name in the table
               self.savedReports.setItem(n, 3, QTableWidgetItem(i[1]))
               self.savedReports.setCursor(Qt.CursorShape.PointingHandCursor)
               self.savedReports.setItem(n, 0, QTableWidgetItem(str(i[0])))
               # self.savedReports.itemChanged.connect(lambda item, report_id=i[0]: self.saveReportName(report_id, item.text()))

               radio_button = QRadioButton(self.windowSaved)
               self.savedReports.setIndexWidget(self.savedReports.model().index(n, 2), radio_button)
               radio_button.setStyleSheet("background-color:white")
               radio_button.setGeometry(124, 31 + (n * 32), 50, 20) 
               radio_button.clicked.connect(lambda _, report_id=i[0]: self.creating(str(report_id)))

            extractAllButton = QPushButton("word تصدير", self.windowSaved)
            extractAllButton.clicked.connect(self.exportAllReportsAsWord)
            extractAllButton.setCursor(Qt.CursorShape.PointingHandCursor)
            extractAllButton.setStyleSheet("background-color:green")
            extractAllButton.setGeometry(320, 460, 160, 30)
            
            extractAllButton = QPushButton("pdf تصدير", self.windowSaved)
            extractAllButton.clicked.connect(self.exportAllReportsAsPdf)
            extractAllButton.setCursor(Qt.CursorShape.PointingHandCursor)
            extractAllButton.setStyleSheet("background-color:green")
            extractAllButton.setGeometry(150, 460, 160, 30)

            self.windowSaved.show()                   
        
        # def saveReportName(self,report_id):
        #   cr.execute("UPDATE reports SET reportName = ? WHERE id = ?", ("new_name", report_id))
        #   con.commit()  


        # Delete report from saved reports
        def deleteReport(self,row,fRom="Original"):

            try:
                if fRom=="OutSide":
                  d = QMessageBox(parent=self.windowCreating,text=f"تأكيد حذف تقرير {self.TableSummary.item(row,6).text()}")
                else:
                  d = QMessageBox(parent=self.windowCreating,text=f"تأكيد حذف تقرير {self.savedReports.item(row,3).text()}")
                d.setIcon(QMessageBox.Icon.Information)
                d.setWindowTitle(title)
                d.setStandardButtons(QMessageBox.StandardButton.Cancel|QMessageBox.StandardButton.Ok)
                important = d.exec()

                if important == QMessageBox.StandardButton.Ok:
    
                    if fRom=="OutSide":
                      cr.execute(f"DELETE FROM reports WHERE id = '{self.TableSummary.item(row,6).text()}'")
                    else:
                      cr.execute(f"DELETE FROM reports WHERE id = '{self.savedReports.item(row,0).text()}'")

                    con.commit()
                    if fRom=="OutSide":
                      self.TableSummary.hideRow(row)
                    else:
                      self.savedReports.hideRow(row)
                    self.load_data()
            except Exception as e:
                print(e)
                d = QMessageBox(parent=self.windowCreating)  
                d.setWindowTitle("فشل")  
                d.setText("حدث خطأ حاول مرة أخرى")
                d.setIcon(QMessageBox.Icon.Warning)
                d.exec() 

        # Function To Create New Report 
        def createReportFun(self):                   

            self.windowCreate = Choices()
            uic.loadUi("design/Create.ui",self.windowCreate)
            self.windowCreate.setFixedSize(359,370)
            self.windowCreate.setWindowTitle(title)
            self.windowCreate.setWindowIcon(QIcon("icons/icon.ico"))
            self.windowCreate.CreateButton.setStyleSheet("background-color: cyan")
            self.windowCreate.CreateButton.clicked.connect(lambda:self.creating("Local"))
            self.windowCreate.setCursor(Qt.CursorShape.PointingHandCursor)
            self.windowCreate.manyPic.addItem("بدون")
            self.windowCreate.manyPic.addItem("1")
            self.windowCreate.manyPic.addItem("2")
            self.windowCreate.manyPic.addItem("3")
            self.windowCreate.manyPic.addItem("4")
            self.windowCreate.manyPic.setCurrentText("4")
            self.windowCreate.NamePrograme.setCheckState(Qt.CheckState.Checked)
            self.windowCreate.Goals.setCheckState(Qt.CheckState.Checked)
            self.windowCreate.Description.setCheckState(Qt.CheckState.Checked)
            self.windowCreate.WhenDate.setCheckState(Qt.CheckState.Checked)
            self.windowCreate.Benefits.setCheckState(Qt.CheckState.Checked)
            self.windowCreate.BenefitsCount.setCheckState(Qt.CheckState.Checked)
            self.windowCreate.Creator.setCheckState(Qt.CheckState.Checked)
            self.windowCreate.show()    

        #Function To Open Control Panel
        def controlPanel(self):

            self.windowControl = Choices()
            self.windowControl.setFixedSize(300,570)
            self.windowControl.setWindowTitle(title)
            self.windowControl.setWindowIcon(QIcon("icons/icon.ico"))
            self.picPathMinLogo = ""
            self.picBinaryMinLogo = ""
            layout = QVBoxLayout()

            self.windowControl.setLayout(layout)
            self.lineone = QLineEdit(self.windowControl)
            Label1 = QLabel("الترويسة الأولى",self.windowControl)
            Label1.setFont(QFont("Normal",15))
            Label1.move(198,10)
            self.lineone.setGeometry(7,40,290,30)

            self.linetwo = QLineEdit(self.windowControl)
            Label2 = QLabel("الترويسة الثانية",self.windowControl)
            Label2.setFont(QFont("Normal",15))
            Label2.move(198,70)
            self.linetwo.setGeometry(7,100,290,30)

            self.linethree = QLineEdit(self.windowControl)
            Label3 = QLabel("الترويسة الثالثة",self.windowControl)
            Label3.setFont(QFont("Normal",15))
            Label3.move(198,130)
            self.linethree.setGeometry(7,160,290,30)

            self.linefour = QLineEdit(self.windowControl)
            Label4 = QLabel("الترويسة الرابعة",self.windowControl)
            Label4.setFont(QFont("Normal",15))
            Label4.move(198,190)
            self.linefour.setGeometry(7,220,290,30)

            Label5 = QLabel("شعار وزارة التعليم",self.windowControl)
            Label5.setGeometry(90,250,120,30)

            self.FrameMin = QFrame(self.windowControl)
            self.FrameMin.setStyleSheet("background-color:rgb(178,178,178)")
            self.FrameMin.setGeometry(30,290,250,200)

            self.layoutFrameLogo = QVBoxLayout()
            self.FrameMin.setLayout(self.layoutFrameLogo)

            picLabel = QLabel(self.FrameMin)
            
            try:
                cr.execute("SELECT icon FROM start")

                with open("images/logo.png","wb") as logoImpo:
                   logoImpo.write(cr.fetchone()[0])
                img = Image.open("images/logo.png")
                img = img.resize((220,125),Image.LANCZOS)
                img.save("images/logo.png",quality=100)
                pix = QPixmap("images/logo.png")
                picLabel.setPixmap(pix)
            except:
                pass

            self.layoutFrameLogo.addWidget(picLabel)

            specialButtonS = QPushButton(self.windowControl)
            specialButtonS.setStyleSheet(f"Qproperty-icon:url(images/cam.png);qproperty-iconSize:15px 15px;background-color:rgb(253, 253, 253)")
            specialButtonS.clicked.connect(self.addPicLogo)
            specialButtonS.setCursor(Qt.CursorShape.PointingHandCursor)
            specialButtonS.setGeometry(130,490,30,30)

            self.SaveButton = QPushButton("حفظ",self.windowControl)
            self.SaveButton.setCursor(Qt.CursorShape.PointingHandCursor)
            self.SaveButton.setGeometry(55,530,200,30)


            cr.execute("SELECT * from start")
            values = cr.fetchall()[0]

            self.SaveButton.clicked.connect(self.Save)
            self.lineone.setText(values[0])
            self.linetwo.setText(values[1])
            self.linethree.setText(values[2])
            self.linefour.setText(values[3])
            self.windowControl.show()

        def addPicLogo(self):

            responce = QFileDialog.getOpenFileName(self.windowControl,"اختر ملفا",desktopPath,filter="Image File (*.*)")
            if len(responce[0])!=0:
                image = Image.open(responce[0])
                self.picPathMinLogo = responce[0]

                finalImage = image.resize((250,170))
                finalImage.save("logo_image.png",quality=100)

                with open("logo_image.png","rb") as temp_binary:
                    binaryCode12 = temp_binary.read()

                self.picBinaryMinLogo = binaryCode12
                picLabel = QLabel(self.FrameMin)
                pix = QPixmap("logo_image.png")
                picLabel.setPixmap(pix)
                
                for i in reversed(range(self.layoutFrameLogo.count())): 
                    self.layoutFrameLogo.itemAt(i).widget().setParent(None)

                self.layoutFrameLogo.addWidget(picLabel)
                os.remove("logo_image.png")

        # Save control panal info
        def Save(self):

            if self.picBinaryMinLogo !="":
                cr.execute(f"UPDATE start set line1='{self.lineone.text()}' ,line2='{self.linetwo.text()}' , line3='{self.linethree.text()}' , line4 = '{self.linefour.text()}',icon=?",([self.picBinaryMinLogo]))
            else:
                cr.execute(f"UPDATE start set line1='{self.lineone.text()}' ,line2='{self.linetwo.text()}' , line3='{self.linethree.text()}' , line4 = '{self.linefour.text()}'")
            con.commit()
            d = QMessageBox(parent=self.windowCreating,text="تم التعديل بنجاح")
            d.setWindowTitle("نجاح")
            d.setIcon(QMessageBox.Icon.Information)
            d.exec()
            self.windowControl.destroy()
            # the next code is under revision
            os.execv(sys.executable, ['python'] + sys.argv)
            app.closeAllWindows()
            
        # Create new content
        def creating(self,fromW):

            try:
                self.destroy()
                self.close()
                self.windowCreating.destroy()
                self.windowCreating.close()
            except:
                pass
            self.ablePrograme= False
            self.ableGoals= False
            self.ableDescription= False
            self.ableCreator= False
            self.ableDate= False
            self.ableBenefits= False
            self.ableCount= False
            self.countPic = 0
            self.secretLittleThing = ""
            if fromW=="Local":
                self.windowCreate.destroy()
            global yForImpo
            global xForImpo
            yForImpo = 50
            xForImpo = 0
            self.pictersPaths = ["","","",""]
            self.windowCreating = Choices()
            self.windowCreating.setMinimumSize(900,1150)
            self.setMinimumSize(930,500)
            self.showMaximized()
            self.setWindowTitle(title)
            self.windowCreating.setWindowTitle(title)
            self.windowCreating.setWindowIcon(QIcon("icons/icon.ico"))
            button_list_widget = QWidget(self.windowCreating)
            button_list_layout = QVBoxLayout()            
            button_list_widget.setLayout(button_list_layout)
            button_list_widget.setMaximumWidth(200)   
            # Add Buttons Function When Creating New Report
            def addButton(icon, text, callback):
                button = QPushButton(text)     
                button.setStyleSheet(
                                        f"""
                                        font-size: 14px;
                                        qproperty-icon: url('{icon}');
                                        qproperty-iconSize: 25px 25px;
                                        background-color: #2ABCB5;                                        
                                        """
                                    )
                button.setLayoutDirection(Qt.LayoutDirection.RightToLeft)
                button.setMinimumHeight(40) 
                button.setCursor(Qt.CursorShape.PointingHandCursor)
                button.clicked.connect(callback)
                button_list_layout.addWidget(button)

            # Add buttons to the layout
            addButton("images/addNewreport.png", "إنشاء تقرير", self.createReportFun)
            addButton("images/savedReports.png", "التقارير المحفوظة", self.savedReportsFun)
            addButton("images/save.png", "حفظ باسم", self.SavePrograme)
            if fromW != "Local":
                addButton("images/save.png", "حفظ", self.updateAReport)
            addButton("images/clearFields.png", "إفراغ الحقول", self.emptyFieldsFun)
            addButton("images/summary.png", "ملخص التقارير", self.summaryReports)
            addButton("images/printer.png", "الطباعة", self.printDoc)
            addButton("images/control.png", "لوحة التحكم", self.controlPanel)
            addButton("images/dbExport.png", "تصدير  قاعدة البيانات", self.exportDb)
            addButton("images/importDb.png", "إستيراد قاعدة البيانات", self.importDb)
            addButton("images/word.png", "word حفظ بصيغة", self.writeWord)
            addButton("images/pdfIcon.png", "pdf حفظ بصيغة", lambda: self.writePdf(fromWhere="convert"))
            
            companyLogobutton = QPushButton()                
            companyLogobutton.setStyleSheet(f"font-size:14px;qproperty-icon:url('images/companyLogo.png');qproperty-iconSize:150px 100px;background-color:transparent;")   
            companyLogobutton.clicked.connect(self.openwebSite)   
            companyLogobutton.setCursor(Qt.CursorShape.PointingHandCursor)      
            button_list_layout.addWidget(companyLogobutton) 
            report_widget = QWidget(self.windowCreating)
            
            
            # Create a vertical layout to hold the header label and the list widget
            report_widgetlayout = QVBoxLayout()                        
            report_widget.setLayout(report_widgetlayout)  
            report_widget.setMinimumHeight(1150)            
            report_widget.setMinimumSize(920,1150)
            self.hiderFrameshow = QFrame(report_widget)
            self.hiderFrameshow.setStyleSheet("background-color: white")
            self.hidderFramePicshow = ClickableQFrame(self.hiderFrameshow)
            self.hidderFramePicshow.setStyleSheet(f"background-color:#EBEAE9;")
            self.hidderFramePicshow.setGeometry(40,5,250,130)
            self.hidderlayoutPicshow = QVBoxLayout()
            self.hidderFramePicshow.setLayout(self.hidderlayoutPicshow)                                    
            self.hidderFramePicshow.clicked.connect(lambda: self.putImage(f"ReportCover"))
            self.hidderFramePicshow.setCursor(Qt.CursorShape.PointingHandCursor)
            specialButton = QPushButton(self.hiderFrameshow)
            specialButton.setIcon(QIcon("images/cam.png"))
            specialButton.setCursor(Qt.CursorShape.PointingHandCursor)
            specialButton.clicked.connect(lambda:self.putImage(f"ReportCover"))
            
            logoLabel = QLabel(self.hiderFrameshow)
            logoLabel.move(350,10)
            
            try:
                cr.execute("SELECT icon FROM start")

                with open("images/logo.png","wb") as logoImpo:
                   logoImpo.write(cr.fetchone()[0])
                img = Image.open("images/logo.png")
                img = img.resize((220,125),Image.LANCZOS)
                img.save("images/logo.png",quality=100)
                pix = QPixmap("images/logo.png")
                logoLabel.setPixmap(pix)
                
            except:
                pass

            specialButton.move(290,65)
            DeleteButtonHidder = QPushButton(self.hiderFrameshow)
            DeleteButtonHidder.setStyleSheet(f"Qproperty-icon:url(images/trashicon.png);qproperty-iconSize:15px 15px;background-color:rgb(253, 253, 253)")
            DeleteButtonHidder.clicked.connect(lambda:self.deleteImagesTemp(f"ReportCover"))
            DeleteButtonHidder.setCursor(Qt.CursorShape.PointingHandCursor)
            DeleteButtonHidder.move(290,90)

            Frame_text = QFrame(self.hiderFrameshow)
            Frame_text.setStyleSheet("background-color: white")
            Frame_text.setGeometry(655,0,250,140)
            
            text_layout = QVBoxLayout(Frame_text)

            Frame_text.setLayout(text_layout)
            cr.execute("SELECT line1 FROM start")
            Label1 = QLabel(cr.fetchone()[0])
            text_layout.addWidget(Label1)
            
            cr.execute("SELECT line2 FROM start")
            Label2 = QLabel("   "+cr.fetchone()[0])
            text_layout.addWidget(Label2)

            cr.execute("SELECT line3 FROM start")
            Label3 = QLabel(cr.fetchone()[0])
            text_layout.addWidget(Label3)

            cr.execute("SELECT line4 FROM start")
            Label4 = QLabel(cr.fetchone()[0])
            text_layout.addWidget(Label4)

            logoLabel = QLabel(self.hiderFrameshow)
            logoLabel.move(350,10)
            
            try:
                cr.execute("SELECT icon FROM start")
                with open("images/logo.png","wb") as logoImpo:
                  logoImpo.write(cr.fetchone()[0])
                img = Image.open("images/logo.png")
                img = img.resize((220,125),Image.LANCZOS)
                img.save("images/logo.png",quality=100)
                pix = QPixmap("images/logo.png")
                logoLabel.setPixmap(pix)
            except:
                pass

            self.hiderFrameshow.setGeometry(0,3,900,140)
            self.cFrameshow = QFrame(report_widget) 
            self.cFrameshow.setStyleSheet("background-color: white")

            labelGood = QLabel("توثيق برنامج",self.cFrameshow)
            labelGood.setStyleSheet("font-size:20px")
            labelGood.move(410,20)
            
            self.programeNameShow = fromW
            if fromW=="Local":
                self.programeNameShow = "توثيق برنامج"
            if fromW!="Local":
                numberOfPictures = -1
                if self.programeNameShow !="":
                    cr.execute(f"SELECT * FROM reports WHERE id = '{self.programeNameShow}'")
                    listImportant = cr.fetchall()[0]
                    
                    if listImportant[2]!="":
                        self.ablePrograme = True
                        self.createNamePrograme()
                        if listImportant[2] == " ":
                            self.programeNameE.setText(str(listImportant[2]).strip())
                        else:
                            self.programeNameE.setText(str(listImportant[2]))
                    if listImportant[3]!="":
                        self.ableGoals = True
                        self.createGoals()
                        self.programeGoalsE.setText(listImportant[3])
                        if listImportant[3] == " ":
                            self.programeGoalsE.setText(str(listImportant[3]).strip())
                        else:
                            self.programeGoalsE.setText(str(listImportant[3]))

                    if listImportant[4]!="":
                        self.ableDescription = True
                        self.createDescription()
                        self.programeDescriptionE.setText(listImportant[4])
                        if listImportant[4] == " ":
                            self.programeDescriptionE.setText(str(listImportant[4]).strip())
                        else:
                            self.programeDescriptionE.setText(str(listImportant[4]))

                    if listImportant[5]!="":
                        self.ableCreator = True
                        self.executer()
                        self.programeCreatorE.setText(listImportant[5])
                        if listImportant[5] == " ":
                            self.programeCreatorE.setText(str(listImportant[5]).strip())
                        else:
                            self.programeCreatorE.setText(str(listImportant[5]))
                            
                    if listImportant[6]!="":
                        self.ableDate = True
                        self.executeDate()
                        self.programeWhenDateE.setText(listImportant[6])
                        if listImportant[6] == " ":
                            self.programeWhenDateE.setText(str(listImportant[6]).strip())
                        else:
                            self.programeWhenDateE.setText(str(listImportant[6]))

                    if listImportant[7]!="":
                        self.ableBenefits = True
                        self.Benefits()
                        self.programeBenefitsE.setText(listImportant[7])
                        if listImportant[7] == " ":
                            self.programeBenefitsE.setText(str(listImportant[7]).strip())
                        else:
                            self.programeBenefitsE.setText(str(listImportant[7]))

                    if listImportant[8]!="":
                        self.ableCount = True
                        self.CountBenefits()

                        self.CountBenefitsE.setText(listImportant[8])
                        if listImportant[8] == " ":
                            self.CountBenefitsE.setText(str(listImportant[8]).strip())
                        else:
                            self.CountBenefitsE.setText(str(listImportant[8]))
                            
                    if listImportant[9]!="":
                        numberOfPictures+=1
                    if listImportant[10]!="":
                        numberOfPictures+=1  
                    if listImportant[11]!="":
                        numberOfPictures+=1  
                    if listImportant[12]!="":
                        numberOfPictures+=1
                    if numberOfPictures > -1:
                        self.CreatePic(numberOfPictures+1)
                        self.countPic = numberOfPictures+1
                    try:
                        if listImportant[9]!="" and listImportant[9]!=" ":
                            with open(f"pic1.png","wb") as image:
                               image.write(listImportant[9])
                            self.pictersPaths[0] = "pic1.png"
                            self.putImage("Other",0)

                        if listImportant[10]!="" and listImportant[10]!=" ":
                            with open(f"pic2.png","wb") as image:
                               image.write(listImportant[10])
                            self.pictersPaths[1] = "pic2.png"
                            self.putImage("Other",1)

                        if listImportant[11]!="" and listImportant[12]!=" ":
                            with open(f"pic3.png","wb") as image:
                              image.write(listImportant[11])
                            self.pictersPaths[2] = "pic3.png"
                            self.putImage("Other",2)

                        if listImportant[12]!="" and listImportant[12]!=" ":
                            with open(f"pic4.png","wb") as image:
                               image.write(listImportant[12])
                            self.pictersPaths[3] = "pic4.png"
                            self.putImage("Other",3)

                        if listImportant[13]!="":
                            self.secretLittleThing =listImportant[13]
                            self.picLogoBinary = listImportant[13]
                            with open(f"secretThing.png","wb") as image:
                              image.write(self.secretLittleThing)
                            image = Image.open("secretThing.png")
                            finalImage = image.resize((350,200))

                            finalImage.save("secretThing.png",quality=100)
                            self.secretLittleThing = "secretThing.png"
                            self.putImage("Other",1000)
                            self.addAdmins(self.cFrameshow)
                    except:
                        pass
                    # this may cause problems in future
                    self.addAdmins(self.cFrameshow)
                    

                    if listImportant[14]!="":
                        self.label1Maye.setText(listImportant[14])
                    if listImportant[15]!="":
                        self.label2Maye.setText(listImportant[15])

                    if listImportant[16]!="":
                        self.MangerName.setText(listImportant[16])
                    if listImportant[17]!="":
                        self.consultName.setText(listImportant[17])

            else:
                if self.windowCreate.NamePrograme.isChecked():
                    self.ablePrograme = True
                    self.createNamePrograme()
                if self.windowCreate.Goals.isChecked():
                    self.ableGoals = True
                    self.createGoals()
                if self.windowCreate.Description.isChecked():
                    self.ableDescription = True
                    self.createDescription()
                if self.windowCreate.Creator.isChecked():
                    self.ableCreator = True
                    self.executer()
                if self.windowCreate.WhenDate.isChecked():
                    self.ableDate = True
                    self.executeDate()
                if self.windowCreate.Benefits.isChecked():
                    self.ableBenefits = True
                    self.Benefits()
                if self.windowCreate.BenefitsCount.isChecked():
                    self.ableCount = True
                    self.CountBenefits()
                if self.windowCreate.manyPic.currentText() != "بدون":
                    self.countPic = self.windowCreate.manyPic.currentText()
                    self.CreatePic(self.windowCreate.manyPic.currentText())
                self.addAdmins(self.cFrameshow)
            self.cFrameshow.setGeometry(0,145,900,1150)
            

            # ******************* left top section ******************* left top section  *******************

            layout_widget = QWidget(self.windowCreating)
            layout_widget.setMaximumWidth(330)

            # Create a vertical layout to hold the header label and the list widget
            layout = QVBoxLayout()                        
            layout_widget.setLayout(layout)              

            # Create the header label
            header_label = QLabel("التقارير المحفوظة")
            header_label.setAlignment(Qt.AlignmentFlag.AlignCenter)             
            header_label.setStyleSheet("font-size: 16px; font-weight: bold; padding: 10px; background-color: #2ABCB5; color:white")

            # Add the header label to the layout
            layout.addWidget(header_label)
            self.listWidget = QListWidget(layout_widget)
            layout.addWidget(self.listWidget)    
            self.load_data()        
            self.listWidget.clicked.connect(self.onItemClicked)
            self.listWidget.setCursor(Qt.CursorShape.PointingHandCursor)

            layout_widget.setMinimumWidth(380)
            
            self.scroll = QScrollArea()   
            self.scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOn)
            self.scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOn)
            self.scroll.setWidgetResizable(True)
            self.scroll.setWidget(report_widget)
                           
            Gridlayout_widget = QWidget()
            # Create a grid layout
            Gridlayout = QGridLayout()
            Gridlayout_widget.setLayout(Gridlayout)
            # Add buttons to the layout at row 0, and columns 0, 1, and 2      
            Gridlayout.addWidget(self.scroll,0, 0)
            Gridlayout.addWidget(button_list_widget,0, 1)
            # Align the buttons to expand vertically (row stretch)
            Gridlayout.setRowStretch(0, 1)
            splitter = QSplitter(Qt.Orientation.Horizontal, self.windowCreating )            
            splitter.setHandleWidth(5)
            splitter.setStyleSheet("""
                QSplitter::handle:horizontal {
                    background-color: gray;
                }
            """)
            splitter.addWidget(layout_widget)
            splitter.addWidget(Gridlayout_widget)
            self.setCentralWidget(splitter) 
            self.show()
            
                
        def onItemClicked(self, item):
            """Handle double-click on the QListWidgetItem."""
            report_id = item.data(Qt.ItemDataRole.UserRole)       
            self.creating(str(report_id))
        # Show single report
        def showReport(self, programeNameShow):

            self.windowshow = Choices() 
            self.windowshow.resize(920,1200)
            self.windowshow.setMinimumSize(920,1200)                                            
            self.setWindowTitle(title)
            self.windowshow.setWindowTitle(title)
            self.windowshow.setWindowIcon(QIcon("icons/icon.ico"))
            global yForImpo
            global xForImpo
            yForImpo = 50
            xForImpo = 0
    
            # Create the current design frame on the right
            right_frame = QWidget(self.windowshow)
            right_layout = QVBoxLayout(self.windowshow)
            right_frame.setLayout(right_layout)
            right_frame.setMinimumHeight(1350)
            right_frame.setMinimumWidth(900)
            # Set initial sizes for splitter
            # Set up the current design UI (existing code)
            hiderFrameshow = QFrame(right_frame)
            hiderFrameshow.setStyleSheet("background-color: white")        
            hidderFramePicshow = QFrame(hiderFrameshow)
            hidderFramePicshow.setStyleSheet(f"background-color:#EBEAE9;")
            hidderFramePicshow.setGeometry(40, 5, 250, 130)
            hidderlayoutPicshow = QVBoxLayout()
            hidderFramePicshow.setLayout(hidderlayoutPicshow)
            
            Frame_text = QFrame(hiderFrameshow)
            Frame_text.setStyleSheet("background-color: white")
            Frame_text.setGeometry(655, 0, 250, 140)
            text_layout = QVBoxLayout(Frame_text)
            Frame_text.setLayout(text_layout)
            cr.execute("SELECT line1 FROM start")
            Label1 = QLabel(cr.fetchone()[0])
            text_layout.addWidget(Label1)
            cr.execute("SELECT line2 FROM start")
            Label2 = QLabel("   " + cr.fetchone()[0])
            text_layout.addWidget(Label2)
            cr.execute("SELECT line3 FROM start")
            Label3 = QLabel(cr.fetchone()[0])
            text_layout.addWidget(Label3)
            cr.execute("SELECT line4 FROM start")
            Label4 = QLabel(cr.fetchone()[0])
            text_layout.addWidget(Label4)
            
            logoLabel = QLabel(hiderFrameshow)
            logoLabel.move(350, 10)

            try:
                cr.execute("SELECT icon FROM start")
                with open("images/logo.png","wb") as logoImpo:
                  logoImpo.write(cr.fetchone()[0])
                img = Image.open("images/logo.png")
                img = img.resize((220,125),Image.LANCZOS)
                img.save("images/logo.png",quality=100)
                pix = QPixmap("images/logo.png")
                logoLabel.setPixmap(pix)
            except:
                pass

            hiderFrameshow.setGeometry(0, 3, 900, 140)
            cFrameshow = QFrame(right_frame)
            cFrameshow.setStyleSheet("background-color: white")
            # Create the main layout for the windowCreating widget
            layout = QVBoxLayout(self.windowshow)
            layout.addWidget(right_frame)
            self.windowshow.setLayout(layout)
            
            # Set the central widget and layout
            def createNamePrograme():
                global yForImpo
                programeName = QTextEdit(cFrameshow)
                programeName.setText("اسم البرنامج")
                programeName.setGeometry(0,0,100,35)
                programeName.setStyleSheet("background-color: #2ABCB5")
                programeName.setFont(QFont("Arial",15))
                programeName.setDisabled(True)
                programeName.move(700,yForImpo)
                
                programeNameE = QTextEdit(cFrameshow)
                programeNameE.setGeometry(0,0,565,35)    
                programeNameE.setAlignment(Qt.AlignmentFlag.AlignLeft)
                programeNameE.setLayoutDirection(Qt.LayoutDirection.RightToLeft)
                programeNameE.setFont(QFont("Arial",14))
                programeNameE.move(135,yForImpo)
                

                yForImpo +=35
                if listImportant[2] == " ":
                    programeNameE.setText(str(listImportant[2]).strip())
                else:
                    programeNameE.setText(str(listImportant[2]))

            def createGoals():
                global yForImpo
                programeGoalsEName = QTextEdit(cFrameshow)
                programeGoalsEName.setText("\n\n  الاهداف ")
                programeGoalsEName.setGeometry(0,0,100,110)
                programeGoalsEName.setStyleSheet("background-color: #2ABCB5")
                programeGoalsEName.setFont(QFont("Arial",15))
                programeGoalsEName.setDisabled(True)
                programeGoalsEName.move(700,yForImpo)
                

                programeGoalsE = QTextEdit(cFrameshow)
                programeGoalsE.setGeometry(0,0,565,110)            
                programeGoalsE.setAlignment(Qt.AlignmentFlag.AlignJustify)
                programeGoalsE.setFont(QFont("Arial",15))
                programeGoalsE.move(135,yForImpo) 
                yForImpo+=110
                if listImportant[3] == " ":
                    programeGoalsE.setText(str(listImportant[3]).strip())
                else:
                    programeGoalsE.setText(str(listImportant[3]))

            def createDescription():
                global yForImpo
                programeDescriptionEName = QTextEdit(cFrameshow)
                programeDescriptionEName.setText("\n\n  الوصف ")
                programeDescriptionEName.setGeometry(0,0,100,110)
                programeDescriptionEName.setStyleSheet("background-color: #2ABCB5")
                programeDescriptionEName.setFont(QFont("Arial",15))
                programeDescriptionEName.setDisabled(True)
                programeDescriptionEName.move(700,yForImpo)
                
                programeDescriptionE = QTextEdit(cFrameshow)
                programeDescriptionE.setGeometry(0,0,565,110)
                programeDescriptionE.setAlignment(Qt.AlignmentFlag.AlignLeft)
                programeDescriptionE.setFont(QFont("Arial",15))
                programeDescriptionE.move(135,yForImpo)
                if listImportant[4] == " ":
                    programeDescriptionE.setText(str(listImportant[4]).strip())
                else:
                    programeDescriptionE.setText(str(listImportant[4]))
                yForImpo+=110

            def executer():
                global yForImpo
                programeCreatorEName = QTextEdit(cFrameshow)
                programeCreatorEName.setText("المنفذ")
                programeCreatorEName.setGeometry(0,0,100,35)
                programeCreatorEName.setStyleSheet("background-color: #2ABCB5")
                programeCreatorEName.setFont(QFont("Arial",15))
                programeCreatorEName.setDisabled(True)
                programeCreatorEName.move(700,yForImpo)
                

                programeCreatorE = QTextEdit(cFrameshow)
                programeCreatorE.setGeometry(10,10,565,35)
                programeCreatorE.setAlignment(Qt.AlignmentFlag.AlignLeft)
                programeCreatorE.setFont(QFont("Arial",15))
                programeCreatorE.move(135,yForImpo)
                if listImportant[5] == " ":
                    programeCreatorE.setText(str(listImportant[5]).strip())
                else:
                    programeCreatorE.setText(str(listImportant[5]))                
                yForImpo+=35

            def executeDate():
                global yForImpo
                programeWhenDateEName = QTextEdit(cFrameshow)
                programeWhenDateEName.setText("تاريخ التنفيذ")
                programeWhenDateEName.setGeometry(0,0,100,35)
                programeWhenDateEName.setStyleSheet("background-color: #2ABCB5")
                programeWhenDateEName.setFont(QFont("Arial",15))
                programeWhenDateEName.setDisabled(True)
                programeWhenDateEName.move(700,yForImpo)
                
                programeWhenDateE = QTextEdit(cFrameshow)
                programeWhenDateE.setGeometry(10,10,565,35)
                programeWhenDateE.setAlignment(Qt.AlignmentFlag.AlignLeft)
                programeWhenDateE.setFont(QFont("Arial",14))
                programeWhenDateE.move(135,yForImpo)
                yForImpo+=35
                if listImportant[6] == " ":
                    programeWhenDateE.setText(str(listImportant[6]).strip())
                else:
                    programeWhenDateE.setText(str(listImportant[6]))         

            def Benefits():
                global yForImpo
                global xForImpo
                programeBenefitsEName = QTextEdit(cFrameshow)
                programeBenefitsEName.setText("المستفيدون")
                programeBenefitsEName.setGeometry(0,0,100,30)
                programeBenefitsEName.setStyleSheet("background-color: #2ABCB5")
                programeBenefitsEName.setFont(QFont("Arial",13))
                programeBenefitsEName.setDisabled(True)
                programeBenefitsEName.move(700,yForImpo)
            
                programeBenefitsE = QTextEdit(cFrameshow)
                programeBenefitsE.setGeometry(10,10,565,30)            
                programeBenefitsE.setAlignment(Qt.AlignmentFlag.AlignLeft)
                programeBenefitsE.setFont(QFont("Arial",13))
                programeBenefitsE.move(135,yForImpo)
                yForImpo += 30
                if listImportant[7] == " ":
                    programeBenefitsE.setText(str(listImportant[7]).strip())
                else:
                    programeBenefitsE.setText(str(listImportant[7]))        

            def CountBenefits():
                global yForImpo
                global xForImpo
                CountBenefitsEName = QTextEdit(cFrameshow)
                CountBenefitsEName.setText("عدد المستفيدين")
                CountBenefitsEName.setGeometry(0,0,100,30)
                CountBenefitsEName.setStyleSheet("background-color: #2ABCB5")
                CountBenefitsEName.setFont(QFont("Arial",13))
                CountBenefitsEName.setDisabled(True)
                CountBenefitsEName.move(700,yForImpo)
                
                CountBenefitsE = QTextEdit(cFrameshow)
                CountBenefitsE.setGeometry(10,10,565,30)
                CountBenefitsE.setAlignment(Qt.AlignmentFlag.AlignLeft)
                CountBenefitsE.setFont(QFont("Arial",13))
                CountBenefitsE.move(135,yForImpo)   
                if listImportant[8] == " ":
                    CountBenefitsE.setText(str(listImportant[8]).strip())
                else:
                    CountBenefitsE.setText(str(listImportant[8]))                         
            
            picters = []
            layouts = []
            def CreatePic(Count):
                x = 100
                y = 440                
                for i in range(int(Count)):
                    if i > 1:
                        if i==2:
                            x = 100
                            y+=230
                        else:
                            x+=360
                    picters.append(QFrame(cFrameshow))
                    picters[i].setGeometry(0,0,350,200)
                    picters[i].setStyleSheet(f"background-color:#EBEAE9;")
                    picters[i].move(x,y)
                            
                    CreatePictemplayout = QVBoxLayout()
                    picters[i].setLayout(CreatePictemplayout)
                    layouts.append(CreatePictemplayout)                    
                    if i<=1:
                        x+=360

            def putImage(number=-1):                                                              
                    if number==0:
                        cr.execute(f"SELECT pic1 FROM reports WHERE id = '{programeNameShow}'")
                    if number==1:
                        cr.execute(f"SELECT pic2 FROM reports WHERE id = '{programeNameShow}'")
                    if number==2:
                        cr.execute(f"SELECT pic3 FROM reports WHERE id = '{programeNameShow}'")
                    if number==3:
                        cr.execute(f"SELECT pic4 FROM reports WHERE id = '{programeNameShow}'")
                    if number==1000:
                        cr.execute(f"SELECT picLogo FROM reports WHERE id = '{programeNameShow}'")
  
                    with open(f"pic11.png","wb") as image:
                        image.write(cr.fetchone()[0])
                    image = Image.open("pic11.png")
                    if number==1000:
                        finalImage = image.resize((240,110))
                    else:
                        finalImage = image.resize((350,180))

                    finalImage.save("image1.png",quality=100)
                    if number==0:
                        picLabel = QLabel(picters[number])
                        pix = QPixmap("image1.png")
                        picLabel.setPixmap(pix)
                        picLabel.setCursor(Qt.CursorShape.PointingHandCursor)
                        for i in reversed(range(layouts[number].count())): 
                            layouts[number].itemAt(i).widget().setParent(None)
                        layouts[number].addWidget(picLabel)
                    elif number==1:
                        picLabel = QLabel(picters[number])
                        pix = QPixmap("image1.png")
                        picLabel.setPixmap(pix)
                        picLabel.setCursor(Qt.CursorShape.PointingHandCursor)
                        for i in reversed(range(layouts[number].count())): 
                            layouts[number].itemAt(i).widget().setParent(None)
                        layouts[number].addWidget(picLabel)
                    elif number==2:
                        picLabel = QLabel(picters[number])
                        pix = QPixmap("image1.png")
                        picLabel.setPixmap(pix)
                        picLabel.setCursor(Qt.CursorShape.PointingHandCursor)
                        for i in reversed(range(layouts[number].count())): 
                            layouts[number].itemAt(i).widget().setParent(None)
                        layouts[number].addWidget(picLabel)
                    elif number==3:
                        picLabel = QLabel(hidderFramePicshow)
                        pix = QPixmap("image1.png")
                        picLabel.setPixmap(pix)
                        picLabel.setCursor(Qt.CursorShape.PointingHandCursor)
                        for i in reversed(range(layouts[number].count())): 
                            layouts[number].itemAt(i).widget().setParent(None)
                        layouts[number].addWidget(picLabel)
                    elif number==1000:
                        picLabel = QLabel(hidderFramePicshow)
                        pix = QPixmap("image1.png")
                        picLabel.setPixmap(pix)
                        picLabel.setCursor(Qt.CursorShape.PointingHandCursor)
                        for i in reversed(range(hidderlayoutPicshow.count())): 
                            hidderlayoutPicshow.itemAt(i).widget().setParent(None)
                        hidderlayoutPicshow.addWidget(picLabel)

 
            numberOfPictures = -1                
            if programeNameShow !="":               
                cr.execute(f"SELECT * FROM reports WHERE id = '{programeNameShow}'")
                listImportant = cr.fetchall()[0]
                
                if listImportant[2]!="":
                    createNamePrograme()
                if listImportant[3]!="":                    
                    createGoals()
                if listImportant[4]!="":
                    createDescription()
                if listImportant[5]!="":                    
                    executer()
                if listImportant[6]!="":                    
                    executeDate()
                if listImportant[7]!="":                    
                    Benefits()
                if listImportant[8]!="":                    
                    CountBenefits()
                if listImportant[9]!="":
                    numberOfPictures+=1
                if listImportant[10]!="":
                    numberOfPictures+=1  
                if listImportant[11]!="":
                    numberOfPictures+=1  
                if listImportant[12]!="":
                    numberOfPictures+=1
                if numberOfPictures > -1:                    
                    CreatePic(numberOfPictures+1)

                if listImportant[9]!="" and listImportant[9]!=" ":                    
                    with open(f"pic11.png","wb") as image1:
                        image1.write(listImportant[9])
                    putImage(0)

                if listImportant[10]!="" and listImportant[10]!=" ":
                    with open(f"pic22.png","wb") as image2:
                        image2.write(listImportant[10])                    
                    putImage(1)

                if listImportant[11]!="" and listImportant[12]!=" ":
                    with open(f"pic33.png","wb") as image3:
                        image3.write(listImportant[11])
                    putImage(2)

                if listImportant[12]!="" and listImportant[12]!=" ":
                    with open(f"pic44.png","wb") as image4:
                        image4.write(listImportant[12])
                    putImage(3)

                if listImportant[13]!="":                                       
                    with open(f"secretThing1.png","wb") as image5:
                        image5.write(listImportant[13])                    
                    putImage(1000)
                    
                label1Maye = QLineEdit(cFrameshow)
                label1Maye.setGeometry(40,900,180,25)

                consultName = QLineEdit(cFrameshow)
                consultName.setGeometry(40,930,180,25)

                label2Maye = QLineEdit(cFrameshow)
                label2Maye.setGeometry(680,900,180,25)

                MangerName = QLineEdit(cFrameshow)
                MangerName.setGeometry(680,930,180,25)   
                
                if listImportant[14]!="":
                    label1Maye.setText(listImportant[14])
                if listImportant[15]!="":
                    label2Maye.setText(listImportant[15])

                if listImportant[16]!="":
                    MangerName.setText(listImportant[16])
                if listImportant[17]!="":
                    consultName.setText(listImportant[17])
        
            cFrameshow.setGeometry(0,145,900,1150)
            
            scroll = QScrollArea(self.windowshow)
            scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOn)
            scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            scroll.setWidgetResizable(True)            
            scroll.setWidget(right_frame)
            layout.addWidget(scroll)                                            
            self.windowshow.show()

        # Function To Generate Summary For Reports
        def summaryReports(self):

            self.windowsummary = Choices()
            self.windowsummary.resize(720,500)
            self.windowsummary.setWindowTitle(title)
            self.windowsummary.setWindowIcon(QIcon("icons/icon.ico"))
            self.TableSummary = QTableWidget(self.windowsummary)
            self.TableSummary.setColumnCount(7)
            self.TableSummary.setHorizontalHeaderLabels(["","عدد المستفيدين","المستفيدون","تاريخ التنفيذ","المنفذ","اسم البرنامج",""])
            self.TableSummary.setColumnHidden(6,True)
            self.windowsummary.resizeEvent = self.resizeSummary
            self.TableSummary.setColumnWidth(0,200)
            cr.execute("SELECT id,name,executer,executeDate,benefits,countBenefits From reports")
            result = cr.fetchall()
            for r,i in enumerate(result):
                self.TableSummary.insertRow(self.TableSummary.rowCount())
                i = list(i)
                i.insert(6,"")
                for col,c in enumerate(reversed(i)):
                    if col == 0:
                        button = QPushButton()
                        button.setStyleSheet(f"Qproperty-icon:url(images/trashIcon.png);qproperty-iconSize:30px 30px;background-color:rgb(253, 253, 253)")
                        button.clicked.connect(lambda x,row=r:self.deleteReport(row,"OutSide"))
                        button.setCursor(Qt.CursorShape.PointingHandCursor)
                        self.TableSummary.setIndexWidget(self.TableSummary.model().index(r,0),button)
                    else:
                        item = QTableWidgetItem(str(c))
                        self.TableSummary.setItem(r,col,item)
                        
            delegate = LineEditDelegate()
            self.TableSummary.setItemDelegate(delegate)
            self.TableSummary.setGeometry(0,0,720,470)
            self.TableSummary.cellDoubleClicked.connect(self.zoomSumarry)
            
            self.pdfExport = QPushButton("Pdf تصدير",self.windowsummary,clicked=self.exportSummaryAsPdf)
            self.pdfExport.setStyleSheet("background-color:red;font-size:20px")
            self.pdfExport.setCursor(Qt.CursorShape.PointingHandCursor)
            self.pdfExport.setGeometry(0,470,360,30)
            self.WordExport = QPushButton("Word تصدير",self.windowsummary,clicked=self.exportSummaryAsWord)
            self.WordExport.setCursor(Qt.CursorShape.PointingHandCursor)
            self.WordExport.setStyleSheet("background-color:blue;font-size:20px")
            self.WordExport.setGeometry(self.pdfExport.width(),470,360,30)
            self.priwidth = 720
            self.prihei = 500
            self.windowsummary.show()

        def zoomSumarry(self,row,col):

            if col !=0 and col!=8 :
                self.windowsummaryZoom = Choices()
                self.windowsummaryZoom.setFixedSize(400,300)
                self.windowsummaryZoom.setWindowTitle(title)
                self.windowsummaryZoom.setWindowIcon(QIcon("icons/icon.ico"))
                summaryZoomText = QTextEdit(self.windowsummaryZoom)
                summaryZoomText.setGeometry(0,0,400,300)
                summaryZoomText.setAlignment(Qt.AlignmentFlag.AlignLeft)
                summaryZoomText.setFont(QFont("Arial",15))
                summaryZoomText.setText(self.TableSummary.item(row,col).text())
                self.windowsummaryZoom.show()

        def resizeSummary(self,en):

            self.widthChanged = self.windowsummary.width() - self.priwidth
            self.heightChanged = self.windowsummary.height() - self.prihei - 10
            self.widthChanged-5
            self.TableSummary.resize(self.TableSummary.width()+self.widthChanged,self.TableSummary.height()+self.heightChanged)
            self.pdfExport.resize(self.pdfExport.width()+(round(self.widthChanged/2)),self.pdfExport.height())
            self.pdfExport.move(0,self.TableSummary.height()+3)
            self.WordExport.resize(self.WordExport.width()+(round(self.widthChanged/2)),self.WordExport.height())
            self.WordExport.move(self.pdfExport.width(),self.TableSummary.height()+3)

            incread = ((self.TableSummary.width() - 29) // 7) - 4
            if incread >= 87:
                for i in range(self.TableSummary.columnCount()):
                    if i!=0 and i!=8:
                        self.TableSummary.setColumnWidth(i,incread)

            self.priwidth =self.windowsummary.width()
            self.prihei = self.windowsummary.height()- 10
            self.TableSummary.show()

        def updateAReport(self):
            namePrograme = ""
            Goals = ""
            description = ""
            executer = ""
            executeDate = ""
            benefits = ""
            countBenefits = ""
            pic1 = ""
            pic2 = ""
            pic3 = ""
            pic4 = ""
            if self.ablePrograme:
                namePrograme = self.programeNameE.toPlainText() if len(self.programeNameE.toPlainText()) > 0 else " "
            if self.ableGoals:
                Goals = self.programeGoalsE.toPlainText() if len(self.programeGoalsE.toPlainText()) > 0 else " "
            if self.ableDescription:
                description = self.programeDescriptionE.toPlainText() if len(self.programeDescriptionE.toPlainText()) > 0 else " "

            if self.ableCreator:
                executer = self.programeCreatorE.toPlainText() if len(self.programeCreatorE.toPlainText()) > 0 else " "

            if self.ableDate:
                executeDate = self.programeWhenDateE.toPlainText() if len(self.programeWhenDateE.toPlainText()) > 0 else " "

            if self.ableBenefits:
                benefits = self.programeBenefitsE.toPlainText() if len(self.programeBenefitsE.toPlainText()) > 0 else " "

            if self.ableCount:
                countBenefits = self.CountBenefitsE.toPlainText() if len(self.CountBenefitsE.toPlainText()) > 0 else " "

            if self.countPic != 0:
                if self.pictersPaths[0]!="":
                    with open(self.pictersPaths[0],"rb") as binary_image:
                        binaryCode = binary_image.read()
                    pic1 = binaryCode
                else:
                    pic1 = " "
                if self.pictersPaths[1]!="":
                    with open(self.pictersPaths[1],"rb") as binary_image:
                        binaryCode = binary_image.read()
                    pic2 = binaryCode
                else:
                    pic2 = " "
                if self.pictersPaths[2]!="":
                    with open(self.pictersPaths[2],"rb") as binary_image:
                        binaryCode = binary_image.read()
                    pic3 = binaryCode
                else:
                    pic3 = " "
                if self.pictersPaths[3]!="":
                    with open(self.pictersPaths[3],"rb") as binary_image:
                        binaryCode = binary_image.read()
                    pic4 = binaryCode
                else:
                    pic4 = " "
            picLogo = ""
            if self.hidderlayoutPicshow.count() > 0:
                picLogo = self.picLogoBinary
            reportName = self.programeNameShow
            try:
                label1Maybe = self.label1Maye.text()
            except:
                label1Maybe=""
            try:          
               label2Maybe = self.label2Maye.text()
            except:
               label2Maybe=""
            
            try:
                manger = str(self.MangerName.text())
            except:
                manger=""
            
            try:
               co_manger = str(self.consultName.text())
            except:
               co_manger=""
               
            cr.execute(f"""UPDATE reports set name=?,Goals=?,description=?,executer=?,executeDate=?,benefits=?,countBenefits=?,pic1=?,pic2=?,pic3=?,pic4=?,picLogo=?,label1Maybe=?,label2Maybe=?,manger=?,co_manger=? WHERE id = ?""",(namePrograme,Goals,description,executer,executeDate,benefits,countBenefits,pic1,pic2,pic3,pic4,picLogo,label1Maybe,label2Maybe,manger,co_manger,reportName))
            # it doesn't work just this way i have to use it to work
            d = QMessageBox(parent=self.windowCreating,text="تم الحفظ بنجاح")
            d.setWindowTitle("نجاح")
            d.setIcon(QMessageBox.Icon.Information)
            ret = d.exec()
            con.commit()

        # Delete images which used as temp images 
        def deleteImagesTemp(self,neNum):

            if neNum=="ReportCover":
                for i in reversed(range(self.hidderlayoutPicshow.count())): 
                    self.hidderlayoutPicshow.itemAt(i).widget().setParent(None)
                    self.secretLittleThing = ""
            elif neNum==0:
                for i in reversed(range(self.layouts[0].count())): 
                    self.layouts[0].itemAt(i).widget().setParent(None)
                    self.pictersPaths[0] = ""
                    
            elif neNum==1:
                for i in reversed(range(self.layouts[1].count())): 
                    self.layouts[1].itemAt(i).widget().setParent(None)
                    self.pictersPaths[1] = ""

            elif neNum==2:
                for i in reversed(range(self.layouts[2].count())): 
                    self.layouts[2].itemAt(i).widget().setParent(None)
                    self.pictersPaths[2] = ""

            elif neNum==3:
                for i in reversed(range(self.layouts[3].count())): 
                    self.layouts[3].itemAt(i).widget().setParent(None)
                    self.pictersPaths[3] = ""

        def addAdmins(self,frame):

            self.label1Maye = QLineEdit(frame)
            self.label1Maye.setGeometry(40,900,180,25)
            self.consultName = QLineEdit(frame)
            self.consultName.setGeometry(40,930,180,25)
            self.label2Maye = QLineEdit(frame)
            self.label2Maye.setGeometry(680,900,180,25)
            self.MangerName = QLineEdit(frame)
            self.MangerName.setGeometry(680,930,180,25)

        # Save The Report As New File 
        def SavePrograme(self):

            self.saveProgrameWindow = Choices()
            self.saveProgrameWindow.setFixedSize(200,200)
            self.saveProgrameWindow.setWindowTitle(title)
            self.saveProgrameWindow.setWindowIcon(QIcon("icons/icon.ico"))
            self.saveProgrameWindow.setStyleSheet("background-color:white")
            Label = QLabel("اسم الملف",self.saveProgrameWindow)
            Label.move(20,20)
            self.NameEntryProgrameFile = QLineEdit(self.saveProgrameWindow)
            self.NameEntryProgrameFile.setFont(QFont("Arial",15))
            self.NameEntryProgrameFile.move(20,40)
            SaveButton = QPushButton("حفظ",self.saveProgrameWindow,clicked=self.saveReport)
            SaveButton.setCursor(Qt.CursorShape.PointingHandCursor)
            SaveButton.setGeometry(0,0,150,40)
            SaveButton.setStyleSheet("background-color:green")
            SaveButton.setFont(QFont("Arial",15))
            SaveButton.move(20,130)
            self.saveProgrameWindow.show()

        def CreatePic(self,Count):
            x = 100
            y = 440
            self.picters = []
            self.buttons = []
            self.layouts = []
            for i in range(int(Count)):
                if i > 1:
                    if i==2:
                        x = 100
                        y+=230
                    else:
                        x+=360
                self.picters.append(ClickableQFrame(self.cFrameshow))
                self.picters[i].setGeometry(0,0,350,200)
                self.picters[i].setStyleSheet(f"background-color:#EBEAE9;")
                self.picters[i].move(x,y)
                self.picters.append(ClickableQFrame(self.cFrameshow))
                self.picters[i].setCursor(Qt.CursorShape.PointingHandCursor)
                self.picters[i].clicked.connect(lambda x=i: self.putImage(f"{x}"))            
                templayout = QVBoxLayout()
                self.picters[i].setLayout(templayout)
                self.layouts.append(templayout)
                button = QPushButton(self.cFrameshow)
                button.setCursor(Qt.CursorShape.PointingHandCursor)
                tempvar = self.picters[i].geometry()
                buttonx = tempvar.x()
                buttony = tempvar.y()
                self.buttons.append(button)
                self.buttons[i].move(buttonx+tempvar.width()//2 - 10,buttony+tempvar.height())
                self.buttons[i].setIcon(QIcon("images/cam.png"))
                self.buttons[i].setCursor(Qt.CursorShape.PointingHandCursor)
                
                self.buttons[i].setObjectName(f"{i}")
                self.buttons[i].clicked.connect(lambda ch,i=i:self.putImage(f"{i}"))

                DeleteButtonHidderInside = QPushButton(self.cFrameshow)
                DeleteButtonHidderInside.setStyleSheet(f"Qproperty-icon:url(images/trashicon.png);qproperty-iconSize:15px 16px;background-color:rgb(253, 253, 253)")
                DeleteButtonHidderInside.clicked.connect(lambda ch,x=i:self.deleteImagesTemp(x))
                DeleteButtonHidderInside.setCursor(Qt.CursorShape.PointingHandCursor)
                DeleteButtonHidderInside.move(buttonx+tempvar.width()//2 - 35,buttony+tempvar.height())
                if i<=1:
                    x+=360
        # Put image in images container
        def putImage(self,ob,number=-1):
            
            try:
                os.remove("image.png")
            except:
                pass

            if ob=="ReportCover":
                responce = QFileDialog.getOpenFileName(self.windowCreating,"اختر ملفا",desktopPath,filter="Image File (*.*)")
                if len(responce[0])!=0:
                    image = Image.open(responce[0])
                    self.secretLittleThing = responce[0]
                    finalImage = image.resize((240,110))
                    finalImage.save("reportheaderimage.png",quality=100)
                    with open("reportheaderimage.png","rb") as temp_binary:
                        binaryCode12 = temp_binary.read()
                    self.picLogoBinary = binaryCode12
                    picLabel = QLabel(self.hidderFramePicshow)
                    pix = QPixmap("reportheaderimage.png")
                    picLabel.setPixmap(pix)        

                    for i in reversed(range(self.hidderlayoutPicshow.count())): 
                        self.hidderlayoutPicshow.itemAt(i).widget().setParent(None)

                    self.hidderlayoutPicshow.addWidget(picLabel)
                    os.remove("reportheaderimage.png")
            elif ob=="Other":
                if number==0:
                    cr.execute(f"SELECT pic1 FROM reports WHERE id = '{self.programeNameShow}'")
                if number==1:
                    cr.execute(f"SELECT pic2 FROM reports WHERE id = '{self.programeNameShow}'")
                if number==2:
                    cr.execute(f"SELECT pic3 FROM reports WHERE id = '{self.programeNameShow}'")
                if number==3:
                    cr.execute(f"SELECT pic4 FROM reports WHERE id = '{self.programeNameShow}'")
                if number==1000:
                    cr.execute(f"SELECT picLogo FROM reports WHERE id = '{self.programeNameShow}'")
                
                with open(f"pic.png","wb") as image:
                    image.write(cr.fetchone()[0])
                image = Image.open("pic.png")
                if number==1000:
                    finalImage = image.resize((240,110))
                else:
                    finalImage = image.resize((350,180))

                finalImage.save("image.png",quality=100)
                if number==0:
                    picLabel = QLabel(self.picters[number])
                    pix = QPixmap("image.png")
                    picLabel.setPixmap(pix)
                    picLabel.setCursor(Qt.CursorShape.PointingHandCursor)
                    for i in reversed(range(self.layouts[number].count())): 
                        self.layouts[number].itemAt(i).widget().setParent(None)
                    self.layouts[number].addWidget(picLabel)
                elif number==1:
                    picLabel = QLabel(self.picters[number])
                    pix = QPixmap("image.png")
                    picLabel.setPixmap(pix)
                    picLabel.setCursor(Qt.CursorShape.PointingHandCursor)
                    for i in reversed(range(self.layouts[number].count())): 
                        self.layouts[number].itemAt(i).widget().setParent(None)
                    self.layouts[number].addWidget(picLabel)
                elif number==2:
                    picLabel = QLabel(self.picters[number])
                    pix = QPixmap("image.png")
                    picLabel.setPixmap(pix)
                    for i in reversed(range(self.layouts[number].count())): 
                        self.layouts[number].itemAt(i).widget().setParent(None)
                    self.layouts[number].addWidget(picLabel)
                elif number==3:
                    picLabel = QLabel(self.hidderFramePicshow)
                    pix = QPixmap("image.png")
                    picLabel.setPixmap(pix)
                    for i in reversed(range(self.layouts[number].count())): 
                        self.layouts[number].itemAt(i).widget().setParent(None)
                    self.layouts[number].addWidget(picLabel)
                elif number==1000:
                    picLabel = QLabel(self.hidderFramePicshow)
                    pix = QPixmap("image.png")
                    picLabel.setPixmap(pix)
                    for i in reversed(range(self.hidderlayoutPicshow.count())): 
                        self.hidderlayoutPicshow.itemAt(i).widget().setParent(None)
                    self.hidderlayoutPicshow.addWidget(picLabel)

            else:
                responce = QFileDialog.getOpenFileName(self.windowCreating,"اختر ملفا",desktopPath,filter="Image File (*.*)")
                if len(responce[0])!=0:
                    try:
                        os.remove("image.png")
                    except:
                        pass

                    self.pictersPaths[int(ob)]=(responce[0])
                    image = Image.open(responce[0])
                    finalImage = image.resize((350,180))
                    finalImage.save("image.png",quality=100)
                    picLabel = QLabel(self.picters[int(ob)])
                    pix = QPixmap("image.png")
                    picLabel.setPixmap(pix)
                    picLabel.setCursor(Qt.CursorShape.PointingHandCursor)

                    for i in reversed(range(self.layouts[int(ob)].count())): 
                        self.layouts[int(ob)].itemAt(i).widget().setParent(None)
                    self.layouts[int(ob)].addWidget(picLabel)
                    os.remove("image.png")

        def createNamePrograme(self):

            global yForImpo
            programeName = QTextEdit(self.cFrameshow)
            programeName.setText("اسم البرنامج")
            programeName.setGeometry(0,0,100,35)
            programeName.setStyleSheet("background-color: #2ABCB5")
            programeName.setFont(QFont("Arial",15))
            programeName.setDisabled(True)
            programeName.move(700,yForImpo)
            
            self.programeNameE = QTextEdit(self.cFrameshow)
            self.programeNameE.setGeometry(0,0,565,35)            
            self.programeNameE.setAlignment(Qt.AlignmentFlag.AlignLeft)
            self.programeNameE.setFont(QFont("Arial",14))
            self.programeNameE.move(135,yForImpo)            
            yForImpo +=35                     

        def resizedWindow(self):
            newWidth = (self.width() - self.cFrameshow.width())//2
            self.cFrameshow.move(newWidth,self.cFrameshow.y())
            self.hiderFrameshow.move(newWidth,self.hiderFrameshow.y())

        def createGoals(self):
            global yForImpo
            programeName = QTextEdit(self.cFrameshow)
            programeName.setText("\n\n  الاهداف ")
            programeName.setGeometry(0,0,100,110)
            programeName.setAlignment(Qt.AlignmentFlag.AlignRight)
            programeName.setStyleSheet("background-color: #2ABCB5")
            programeName.setFont(QFont("Arial",15))
            programeName.setDisabled(True)
            programeName.move(700,yForImpo)
            
            self.programeGoalsE = QTextEdit(self.cFrameshow)
            self.programeGoalsE.setGeometry(0,0,565,110)            
            self.programeGoalsE.setAlignment(Qt.AlignmentFlag.AlignLeft)
            self.programeGoalsE.setFont(QFont("Arial",15))
            self.programeGoalsE.move(135,yForImpo) # 280            
            yForImpo+=110

        def createDescription(self):
            global yForImpo
            programeName = QTextEdit(self.cFrameshow)
            programeName.setText("\n\n  الوصف ")
            programeName.setGeometry(0,0,100,110)
            programeName.setStyleSheet("background-color: #2ABCB5")
            programeName.setFont(QFont("Arial",15))
            programeName.setDisabled(True)
            programeName.move(700,yForImpo)
        
            self.programeDescriptionE = QTextEdit(self.cFrameshow)
            self.programeDescriptionE.setGeometry(0,0,565,110)
            self.programeDescriptionE.setAlignment(Qt.AlignmentFlag.AlignLeft)
            self.programeDescriptionE.setFont(QFont("Arial",15))
            self.programeDescriptionE.move(135,yForImpo)            
            yForImpo+=110

        def executer(self):
            global yForImpo
            programeName = QTextEdit(self.cFrameshow)
            programeName.setText("المنفذ")
            programeName.setGeometry(0,0,100,35)
            programeName.setStyleSheet("background-color: #2ABCB5")
            programeName.setFont(QFont("Arial",15))
            programeName.setDisabled(True)
            programeName.move(700,yForImpo)
            

            self.programeCreatorE = QTextEdit(self.cFrameshow)
            self.programeCreatorE.setGeometry(10,10,565,35)
            self.programeCreatorE.setAlignment(Qt.AlignmentFlag.AlignLeft)
            self.programeCreatorE.setFont(QFont("Arial",15))
            self.programeCreatorE.move(135,yForImpo)
            yForImpo+=35

        def executeDate(self):
            global yForImpo
            programeName = QTextEdit(self.cFrameshow)
            programeName.setText("تاريخ التنفيذ")
            programeName.setGeometry(0,0,100,35)
            programeName.setStyleSheet("background-color: #2ABCB5")
            programeName.setFont(QFont("Arial",15))
            programeName.setDisabled(True)
            programeName.move(700,yForImpo)
                    
            self.programeWhenDateE = QTextEdit(self.cFrameshow)
            self.programeWhenDateE.setGeometry(10,10,565,35)
            self.programeWhenDateE.setAlignment(Qt.AlignmentFlag.AlignLeft)
            self.programeWhenDateE.setFont(QFont("Arial",14))
            self.programeWhenDateE.move(135,yForImpo)            
            yForImpo+=35

        def Benefits(self):
            global yForImpo
            global xForImpo
            programeName = QTextEdit(self.cFrameshow)
            programeName.setText("المستفيدون")
            programeName.setGeometry(0,0,100,30)
            programeName.setStyleSheet("background-color: #2ABCB5")
            programeName.setFont(QFont("Arial",13))
            programeName.setDisabled(True)
            programeName.move(700,yForImpo)
            

            self.programeBenefitsE = QTextEdit(self.cFrameshow)
            self.programeBenefitsE.setGeometry(10,10,565,30)            
            self.programeBenefitsE.setAlignment(Qt.AlignmentFlag.AlignLeft)
            self.programeBenefitsE.setFont(QFont("Arial",13))
            self.programeBenefitsE.move(135,yForImpo)            
            yForImpo += 30

        # Clear All Fields Text
        def emptyFieldsFun(self):
            d = QMessageBox(parent=self.windowCreating,text="تأكيد افراغ جميع الحقول")
            d.setIcon(QMessageBox.Icon.Information)
            d.setWindowTitle(title)
            d.setStandardButtons(QMessageBox.StandardButton.Cancel|QMessageBox.StandardButton.Ok)
            important = d.exec()
            if important == QMessageBox.StandardButton.Ok:
                try:
                    self.programeNameE.setText("")
                except:
                    pass
                try:
                    self.programeNameE.setText("")
                except:
                    pass
                try:
                    self.programeGoalsE.setText("")
                except:
                    pass
                try:
                    self.programeDescriptionE.setText("")
                except:
                    pass
                try:
                    self.programeCreatorE.setText("")
                except:
                    pass
                try:
                    self.programeWhenDateE.setText("")
                except:
                    pass
                try:
                    self.programeBenefitsE.setText("")
                except:
                    pass
                try:
                    self.CountBenefitsE.setText("")
                except:
                    pass

                try:
                    for i in reversed(range(self.layouts[0].count())): 
                        self.layouts[0].itemAt(i).widget().setParent(None)
                        self.pictersPaths[0] = ""
                except:
                    pass

                try:
                    for i in reversed(range(self.layouts[1].count())): 
                        self.layouts[1].itemAt(i).widget().setParent(None)
                        self.pictersPaths[1] = ""

                except:
                    pass

                try:
                    for i in reversed(range(self.layouts[2].count())): 
                        self.layouts[2].itemAt(i).widget().setParent(None)
                        self.pictersPaths[2] = ""

                except:
                    pass

                try:
                    for i in reversed(range(self.layouts[3].count())): 
                        self.layouts[3].itemAt(i).widget().setParent(None)
                        self.pictersPaths[3] = ""
                except:
                    pass

                try:
                    self.label1Maye.setText("")
                except:
                    pass
                try:
                    self.label2Maye.setText("")
                except:
                    pass

                try:
                    self.consultName.setText("")
                except:
                    pass

                try:
                    self.MangerName.setText("")
                except:
                    pass

        def CountBenefits(self):
            global yForImpo
            global xForImpo
            programeName = QTextEdit(self.cFrameshow)
            programeName.setText("عدد المستفيدين")
            programeName.setGeometry(0,0,100,30)
            programeName.setStyleSheet("background-color: #2ABCB5")
            programeName.setFont(QFont("Arial",13))
            programeName.setDisabled(True)
            programeName.move(700,yForImpo)
            

            self.CountBenefitsE = QTextEdit(self.cFrameshow)
            self.CountBenefitsE .setGeometry(10,10,565,30)
            self.CountBenefitsE .setAlignment(Qt.AlignmentFlag.AlignRight)
            self.CountBenefitsE .setFont(QFont("Arial",13))
            self.CountBenefitsE.move(135,yForImpo)

        def saveReport(self):

            namePrograme = ""
            Goals = ""
            description = ""
            executer = ""
            executeDate = ""
            benefits = ""
            countBenefits = ""
            pic1 = ""
            pic2 = ""
            pic3 = ""
            pic4 = ""
            if self.ablePrograme:
                namePrograme = self.programeNameE.toPlainText() if len(self.programeNameE.toPlainText()) > 0 else " "
            if self.ableGoals:
                Goals = self.programeGoalsE.toPlainText() if len(self.programeGoalsE.toPlainText()) > 0 else " "
            if self.ableDescription:
                description = self.programeDescriptionE.toPlainText() if len(self.programeDescriptionE.toPlainText()) > 0 else " "

            if self.ableCreator:
                executer = self.programeCreatorE.toPlainText() if len(self.programeCreatorE.toPlainText()) > 0 else " "

            if self.ableDate:
                executeDate = self.programeWhenDateE.toPlainText() if len(self.programeWhenDateE.toPlainText()) > 0 else " "

            if self.ableBenefits:
                benefits = self.programeBenefitsE.toPlainText() if len(self.programeBenefitsE.toPlainText()) > 0 else " "

            if self.ableCount:
                countBenefits = self.CountBenefitsE.toPlainText() if len(self.CountBenefitsE.toPlainText()) > 0 else " "
            if self.countPic != 0:
                if self.pictersPaths[0]!="":
                    with open(self.pictersPaths[0],"rb") as binary_image:
                        binaryCode = binary_image.read()
                    pic1 = binaryCode
                else:
                    pic1 = " "
                if self.pictersPaths[1]!="":
                    with open(self.pictersPaths[1],"rb") as binary_image:
                        binaryCode = binary_image.read()
                    pic2 = binaryCode
                else:
                    pic2 = " "
                if self.pictersPaths[2]!="":
                    with open(self.pictersPaths[2],"rb") as binary_image:
                        binaryCode = binary_image.read()
                    pic3 = binaryCode
                else:
                    pic3 = " "
                if self.pictersPaths[3]!="":
                    with open(self.pictersPaths[3],"rb") as binary_image:
                        binaryCode = binary_image.read()
                    pic4 = binaryCode
                else:
                    pic4 = " "
            picLogo = ""
            if self.hidderlayoutPicshow.count() > 0:
                picLogo = self.picLogoBinary
            reportName = str(self.NameEntryProgrameFile.text())
            try:
                label1Maybe = self.label1Maye.text()
            except:
                label1Maybe=''
            
            try:
                label2Maybe = self.label2Maye.text()
            except:
                label2Maybe=''

            try:
                manger = str(self.MangerName.text())
            except:
                manger=''

            try:
                co_manger = str(self.consultName.text())
            except:
                co_manger=''

            cr.execute(f"""Insert INTO reports (reportName,name,Goals,description,executer,executeDate,benefits,countBenefits,pic1,pic2,pic3,pic4,picLogo,label1Maybe,label2Maybe,manger,co_manger) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",(reportName,namePrograme,Goals,description,executer,executeDate,benefits,countBenefits,pic1,pic2,pic3,pic4,picLogo,label1Maybe,label2Maybe,manger,co_manger))
            d = QMessageBox(parent=self.windowCreating,text="تم الحفظ بنجاح")
            d.setWindowTitle("نجاح")
            d.setIcon(QMessageBox.Icon.Information)
            d.exec()
            self.saveProgrameWindow.destroy()
            con.commit()
            self.load_data()  

        # Get Pdf for printer and converter
        def writePdf(self,fromWhere):
            # this line reinit admins fields
            # self.addAdmins(self.cFrameshow)
            content = [] 
            if fromWhere=="convert":
               FileNameSave = QFileDialog.getSaveFileName(self.windowCreating, "اختر مسارا", desktopPath, "PDF Documents (*.pdf);;All Files (*)")
            else:
               file_path = "printFile.pdf"
               FileNameSave = [file_path]
          
            if len(FileNameSave[0]) > 0:
                pdf_file_path = FileNameSave[0]
                doc = SimpleDocTemplate(pdf_file_path, pagesize=letter, rightMargin=0, leftMargin=0, topMargin=30, bottomMargin=5)

                # Register the Amiri font
                font_path = 'font/Amiri-Regular.ttf' 
                pdfmetrics.registerFont(TTFont('ArabicFont', font_path))
                font_path_bold = 'font/Amiri-Bold.ttf'  
                pdfmetrics.registerFont(TTFont('ArabicFont-Bold', font_path_bold))
                # Set up styles
                styles = getSampleStyleSheet()
                custom_style = ParagraphStyle('CustomStyle', parent=styles['Normal'], fontSize=16, spaceAfter=14, alignment=1)
                custom_style.fontName = 'ArabicFont-bold'  

                custom_style_header = ParagraphStyle('CustomStyle', parent=styles['Normal'], fontSize=14, alignment=TA_RIGHT)
                custom_style_header.fontName = 'ArabicFont'  
                

                # # Prepare header data with text and placeholders for images
                # cr.execute("SELECT line1, line2, line3, line4 FROM start")
                # lines = cr.fetchone()
               

                # header_data = [[
                # "",  # Placeholder for the second column (image)
                # "",  # Placeholder for the third column (image)
                # Paragraph(
                # "<br/><br/>".join(
                #    get_display(arabic_reshaper.reshape(line)) for line in lines
                #  ),
                # custom_style_header
                # ),
                #  ]]

                # Prepare header data with text and placeholders for images
                cr.execute("SELECT line1 FROM start")
                line1=cr.fetchone()[0]
                cr.execute("SELECT line2 FROM start")
                line2=cr.fetchone()[0]
                cr.execute("SELECT line3 FROM start")
                line3=cr.fetchone()[0]
                cr.execute("SELECT line4 FROM start")
                line4=cr.fetchone()[0]
                header_data = [[
                 "", 
                 "",  
                Paragraph(f"""
                          {
                            get_display(f"{arabic_reshaper.reshape(line1)}")  
                          }
                          {'&nbsp;' * 4}
                            <br/><br/>
                          {
                              get_display(arabic_reshaper.reshape(line2))
                           }
                           {'&nbsp;' * 8}
                            <br/><br/>
                           {
                              get_display(arabic_reshaper.reshape(line3))
                           }
                            <br/><br/>
                          {
                              get_display(arabic_reshaper.reshape(line4))
                           }
                          """, 
                          custom_style_header
                          ), 
                 ]]
       
                
                # Add images to the header
                if os.path.exists("images/logo.png"):
                  piclogo =img("images/logo.png", width=160, height=80) 
                  header_data[0][1] = piclogo  # Assign the first image to the second column

                if os.path.exists(self.secretLittleThing):
                  logo2 = img(self.secretLittleThing, width=165, height=80)  # Adjust dimensions as needed
                  header_data[0][0] = logo2  # Assign the second image to the third column
                else:
                  pass
          
                # Calculate dynamic column widths
                page_width = letter[0]    # The width of the letter page
                image_width = 2.6 * inch  # Width allocated for each image column
                max_text_width = page_width - (2.3 * image_width) - 1 * inch  # Remaining width for the first column
                # Create the header table
                header_table = Table(header_data, colWidths=[max_text_width,image_width,image_width])
                header_table.setStyle(TableStyle([
                 ('SIZE', (0, 0), (-1, -1), 14), 
                 ('VALIGN', (1, 0), (1, 0), 'RIGHT'),
                 ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),  # Align all content to the right
                 ('RIGHTPADDING', (0, 0), (-1, -1), -25),  # Remove right padding for all cells
                 ('FONTNAME', (0, 0), (-1, 0),'ArabicFont'),
                 ('BOTTOMPADDING', (0, 0), (-1, 0), 16)
                                 
                ]))

                content.append(header_table)
                # Add additional text (Title)
                title = Paragraph(get_display(arabic_reshaper.reshape("\t\t\t\t\t\t\tتوثيق برنامج")), custom_style)
                content.append(title)

               # *Get Report Content*
               # Style for the section names (right column)
                section_right_style = ParagraphStyle(
                  name="rightContent",
                  alignment=TA_RIGHT,
                  textColor=colors.black,
                  fontSize=12,
                  spaceAfter=10,
                  leading=18
                )
                
                section_right_style.fontName = 'ArabicFont-bold'  # Set the custom style font to ArabicFont
                # Style for the section content (left column)
                section_left_style = ParagraphStyle(
                  name="leftContent",
                  alignment=TA_RIGHT,
                  leading=15,  # Adjust this value to set the desired line spacing
                  fontSize=12,
                  spaceAfter=10
                )
                section_left_style.fontName = 'ArabicFont' 

                data=[]
                if self.ablePrograme:
                  data.append(
                    ((get_display(arabic_reshaper.reshape("اسم البرنامج")), get_display(arabic_reshaper.reshape(self.programeNameE.toPlainText()))))
                   )
                if self.ableGoals:
                  goals_text = get_display(arabic_reshaper.reshape(self.programeGoalsE.toPlainText()))
                  # Ensure line breaks are preserved
                  goals_text = goals_text.replace('\n', '<br/>')  # Convert newline characters to <br/>
                  data.append(   
                    (get_display(arabic_reshaper.reshape("الأهداف")),goals_text)
                  )
                if self.ableDescription:
                  description_text = get_display(arabic_reshaper.reshape(self.programeDescriptionE.toPlainText()))
                  # Ensure line breaks are preserved
                  description_text = description_text.replace('\n', '<br/>')  # Convert newline characters to <br/>
                  data.append(   
                    (get_display(arabic_reshaper.reshape("الوصف")),description_text)
                  )
                if self.ableCreator:
                  data.append(   
                   (get_display(arabic_reshaper.reshape("المنفذ")),get_display(arabic_reshaper.reshape(self.programeCreatorE.toPlainText())))
                   ) 
                if self.ableDate:
                   data.append(   
                   (get_display(arabic_reshaper.reshape("تاريخ التنفيذ")),get_display(arabic_reshaper.reshape(self.programeWhenDateE.toPlainText())))
                   )
                if self.ableBenefits:
                   data.append(   
                   (get_display(arabic_reshaper.reshape("المستفيدون")),get_display(arabic_reshaper.reshape(self.programeBenefitsE.toPlainText())))
                   ) 
                if self.ableCount:
                   data.append(   
                   (get_display(arabic_reshaper.reshape("عدد المستفيدين")),get_display(arabic_reshaper.reshape(self.CountBenefitsE.toPlainText())))
                   )
        
                # Create the table data by formatting each section name and content
                table_items = []
                for section_name, section_content in data:
                    # Right column: Section name (red text)
                    right_col = Paragraph(section_name,section_right_style)
                    # Left column: Section content
                    left_col = Paragraph(section_content,section_left_style)
                    # Append the two-column row to table_data
                    table_items.append([left_col, right_col])
                table = Table(table_items, colWidths=[6.6 * inch, 1.3 * inch])

                # Add some basic styling to the table (optional)
                table.setStyle(TableStyle([
                  ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),  # Center vertically
                  ('ALIGN', (0, 0), (-1, -1), 'CENTER'),    # Center horizontally
                  ('BACKGROUND', (1, 0), (1, -1), colors.HexColor("#2ABCB5")),  # Right cell background
                  ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.black),
                  ('BOX', (0, 0), (-1, -1), 0.1, colors.black),
                  ('FONTNAME', (0, 0), (-1, 0), 'ArabicFont'),  # Use the registered font name
                  ('FONTNAME', (0, 1), (-1, -1), 'ArabicFont'),  # Use the registered font name for data
                  ('TOPPADDING', (0, 0), (-1, -1), 3) # Add top padding
                ]))

                content.append(table)

            if self.countPic != 0:
                images_data = []
                max_images_per_row = 2 
                current_row = []

                for i in range(len(self.pictersPaths)):
                    if self.pictersPaths[i] != "":
                        # Calculate dynamic width for images
                        page_width, _ = letter  # Get the width of the page
                        image_width = (page_width / max_images_per_row) * 0.8  # Set image width to 80% of the divided space
                        image = img(self.pictersPaths[i], width=281, height=135)  # Adjust dimensions as needed

                        current_row.append(image)

                        if len(current_row) == max_images_per_row:
                          images_data.append(current_row)  # Add the current row without empty cells
                          current_row = []  # Reset for the next row

                if current_row:
                  images_data.append(current_row)  # Add any remaining images in the last row

                if images_data:
                  # Calculate the total number of columns
                  images_table = Table(images_data, colWidths=[4 * inch, 4 * inch])
                  images_table.setStyle(TableStyle([
                   ('SIZE', (0, 0), (-1, -1), 30),
                   ('FONTNAME', (0, 0), (-1, -1), 'ArabicFont'),
                   ('TOPPADDING', (0, 0), (-1, -1), 2.5),  
                   ('BOTTOMPADDING', (0, 0), (-1, -1), 0),  
                   ('LEFTPADDING', (0, 0), (-1, -1), 3.9  ) 
                   ]))

                  content.append(images_table)

            # *Get Footer Content*
            # Style for  (right column)
            footer_right_style = ParagraphStyle(
                name="rightContent",
                alignment=TA_RIGHT,
                textColor=colors.black,
                fontSize=12,
                spaceAfter=10
            )
            footer_right_style.fontName = 'ArabicFont-bold'  # Set the custom style font to ArabicFont
             # Style for  (left column)
            footer_left_style = ParagraphStyle(
                name="leftContent",
                alignment=TA_LEFT,
                fontSize=12,
                spaceAfter=10
            )
            footer_left_style.fontName = 'ArabicFont-bold' 
            footer_data = [[[],[]],[[],[]]] 
            try:
                if len(self.label2Maye.text()) > 0:
                   footer_data[0][0].append(get_display(arabic_reshaper.reshape(self.label2Maye.text())))
                if len(self.consultName.text()) > 0:
                   footer_data[1][1].append(get_display(arabic_reshaper.reshape(self.consultName.text())))
                if len(self.label1Maye.text()) > 0:
                   footer_data[0][1].append(get_display(arabic_reshaper.reshape(self.label1Maye.text())))
                if len(self.MangerName.text()) > 0:
                   footer_data[1][0].append(get_display(arabic_reshaper.reshape(self.MangerName.text())))

                footer_table_items = []
                for first_list, second_list in footer_data:
                    # Convert lists to strings before creating Paragraphs
                    first = ' '.join(first_list) if first_list else ''
                    second = ' '.join(second_list) if second_list else ''
                    right_col = Paragraph(first, footer_right_style)
                    left_col = Paragraph(second, footer_left_style)
                    # Append the two-column row to table_data
                    footer_table_items.append([left_col, right_col])

                # Create the Table
                table = Table(footer_table_items, colWidths=[3.7 * inch, 3.7 * inch])

                # Add some basic styling to the table (optional)
                table.setStyle(TableStyle([
                  ('FONTNAME', (0, 0), (-1, 0), 'ArabicFont'),  # Use the registered font name
                  ('FONTNAME', (0, 1), (-1, -1), 'ArabicFont'),  # Use the registered font name for data
                  ('LEFTPADDING', (0, 0), (-1, -1), 10),  # Add left padding
                   ('RIGHTPADDING', (0, 0), (-1, -1), 10),  # Add right padding
                   ('TOPPADDING', (0, 0), (-1, -1), 10),  # Add top padding
                  ('BOTTOMPADDING', (0, 0), (-1, -1), 10),  # Add bottom padding
                ]))

                content.append(table)
            except:
                pass
            def add_border(canvas, doc):
                # Draw a border around the page
                width, height = letter  # Get the page dimensions
                border_offset = 20  # Thickness of the border
                canvas.setStrokeColor(colors.black)
                canvas.setLineWidth(1)  # Border width
                canvas.rect(border_offset, border_offset, width - 2 * border_offset, height - 2 * border_offset, stroke=1, fill=0)

            # Build the PDF document
            doc.build(content, onFirstPage=add_border, onLaterPages=add_border) 
            if fromWhere == "convert":
              d = QMessageBox(parent=self.windowCreating,text=f"تم التصدير بنجاح")
              d.setWindowTitle("نجاح")
              d.setIcon(QMessageBox.Icon.Information)
              d.exec()
            else:
                pass
        def writeWord(self):            
            FileNameSave = QFileDialog.getSaveFileName(self.windowCreating,"Select File",desktopPath)
            if len(FileNameSave[0])>0:
                folder = (str(FileNameSave[0]).split(f"/"))
                nameFile = folder[-1]
                folderFinle = "/".join(folder[:-1])
                doc = docx.Document()
                sections = doc.sections
                for section in sections:
                    section.top_margin = docx.shared.Cm(0.7)
                    section.bottom_margin = docx.shared.Cm(0.7)
                    section.left_margin = docx.shared.Cm(0.7)
                    section.right_margin = docx.shared.Cm(0.7)
                sec_pr = doc.sections[0]._sectPr # get the section properties el
                # create new borders el
                pg_borders = OxmlElement('w:pgBorders')
                # specifies how the relative positioning of the borders should be calculated
                pg_borders.set(qn('w:offsetFrom'), 'page')
                for border_name in ('top', 'left', 'bottom', 'right',): # set all borders
                    border_el = OxmlElement(f'w:{border_name}')
                    border_el.set(qn('w:val'), 'triple') # a single line
                    border_el.set(qn('w:sz'), '4') # for meaning of  remaining attrs please look docs
                    border_el.set(qn('w:space'), '10')
                    border_el.set(qn('w:color'), 'black')
                    pg_borders.append(border_el) # register single border to border el
                sec_pr.append(pg_borders) # apply border changes to section

                headers_table = doc.add_table(rows=1, cols=2)
                for row in headers_table.rows:
                    for cell in row.cells:
                        tc = cell._element.tcPr
                        tc.left = None
                        tc.top = None
                        tc.right = None
                        tc.bottom = None
                        cell.vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        cell.paragraphs[0].alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
                        cell.paragraphs[0].size = docx.shared.Pt(15)

                hdr_Cells = headers_table.rows[0].cells

                cr.execute("SELECT line1 FROM start")
                hdr_Cells[1].text = cr.fetchone()[0]
                cr.execute("SELECT line2 FROM start")
                hdr_Cells[1].text = hdr_Cells[1].text+"\n"+cr.fetchone()[0]+"\t"
                cr.execute("SELECT line3 FROM start")
                hdr_Cells[1].text = hdr_Cells[1].text+"\n"+cr.fetchone()[0]
                cr.execute("SELECT line4 FROM start")
                hdr_Cells[1].text = hdr_Cells[1].text+"\n"+cr.fetchone()[0]

                widths = (docx.shared.Inches(5.8),docx.shared.Inches(3))
                for row in headers_table.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width

                heights = (docx.shared.Inches(1.1),docx.shared.Inches(1.1))
                for idx,row in enumerate(headers_table.rows):
                    row.height = heights[idx]

                paragraph12322 =hdr_Cells[1].paragraphs[0]
                run = paragraph12322.runs
                font = run[0].font
                font.size= docx.shared.Pt(15)

                cells = headers_table.rows[0].cells[0].paragraphs[0]
                runCells = cells.add_run()
                if self.secretLittleThing !="":
                    runCells.add_picture(self.secretLittleThing,width=docx.shared.Inches(2.1),height=docx.shared.Inches(1))
                if self.hidderlayoutPicshow.count() <= 0:
                    runCells.add_text("\t\t\t\t\t")
                else:
                    runCells.add_text("\t")
                xsaw = runCells.add_picture("images/logo.png",width=docx.shared.Inches(2.5),height=docx.shared.Inches(1))

                for row in headers_table.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                headers_table.rows[0].cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT


                """
                if self.windowCreate.BenefitsCount.isChecked():
                    self.CountBenefits()
                """

                GoodPrograme = doc.add_paragraph("\t\t\t\t\t\t\tتوثيق برنامج")
                GoodPrograme.runs[0].font.size = docx.shared.Pt(20)
                GoodPrograme.paragraph_format.space_after = docx.shared.Pt(0.1)
                GoodPrograme.paragraph_format.space_before = docx.shared.Pt(1)

                if self.ablePrograme:
                    text = self.programeNameE.toPlainText()
                    listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                    final_text = []
                    for i in text:
                        if i in listNubmers:
                            final_text.append(convert_numbers.english_to_hindi(i))
                        else:
                            final_text.append(i)

                    programeNameProgrameTable = doc.add_table(rows=1,cols=2)

                    programeNameProgrameTable.style = "Table Grid"
                    hdr_Cells = programeNameProgrameTable.rows[0].cells
                    hdr_Cells[1].text = "اسم البرنامج"
                    hdr_Cells[0].text = "".join(final_text)
                    programeNameProgrameTable.autofit = False

                    cell_xml_element = programeNameProgrameTable.rows[0].cells[1]._tc
                    table_cell_properties = cell_xml_element.get_or_add_tcPr()
                    shade_obj = OxmlElement("w:shd")
                    shade_obj.set(qn2("w:fill"),"2ABCB5")
                    table_cell_properties.append(shade_obj)

                    widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))

                    for row in programeNameProgrameTable.rows:
                        for idx, width in enumerate(widths):
                            row.cells[idx].width = width
                    
                    for row in programeNameProgrameTable.rows:
                        for cell in row.cells:
                            cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.size= docx.shared.Pt(17)
                    heights = (docx.shared.Inches(.35), docx.shared.Inches(.35))
                    for idx,row in enumerate(programeNameProgrameTable.rows):
                            row.height = heights[idx]
                if self.ableGoals:
                    text = self.programeGoalsE.toPlainText()
                    listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                    final_text = []
                    for i in text:
                        if i in listNubmers:
                            final_text.append(convert_numbers.english_to_hindi(i))
                        else:
                            final_text.append(i)

                    programeGolasTable = doc.add_table(rows=1,cols=2)

                    programeGolasTable.style = "Table Grid"
                    hdr_Cells = programeGolasTable.rows[0].cells
                    hdr_Cells[1].text = "\t\tالأهداف"
                    hdr_Cells[0].text = "".join(final_text)
                    programeGolasTable.autofit = False

                    cell_xml_element = programeGolasTable.rows[0].cells[1]._tc
                    table_cell_properties = cell_xml_element.get_or_add_tcPr()
                    shade_obj = OxmlElement("w:shd")
                    shade_obj.set(qn2("w:fill"),"2ABCB5")
                    table_cell_properties.append(shade_obj)

                    widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))
                    heights = (docx.shared.Inches(1.4), docx.shared.Inches(1.4))
                    for row in programeGolasTable.rows:
                        for idx, width in enumerate(widths):
                            row.cells[idx].width = width
                    for idx,row in enumerate(programeGolasTable.rows):
                        row.height = heights[idx]
                    

                    for row in programeGolasTable.rows:
                        for cell in row.cells:
                            cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.size= docx.shared.Pt(18)


                if self.ableDescription:
                    text = self.programeDescriptionE.toPlainText()
                    listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                    final_text = []
                    for i in text:
                        if i in listNubmers:
                            final_text.append(convert_numbers.english_to_hindi(i))
                        else:
                            final_text.append(i)

                    programeDescriptionTable = doc.add_table(rows=1,cols=2)

                    programeDescriptionTable.style = "Table Grid"
                    hdr_Cells = programeDescriptionTable.rows[0].cells
                    hdr_Cells[1].text = "\t\tالوصف"
                    hdr_Cells[0].text = "".join(final_text)
                    programeDescriptionTable.autofit = False

                    cell_xml_element = programeDescriptionTable.rows[0].cells[1]._tc
                    table_cell_properties = cell_xml_element.get_or_add_tcPr()
                    shade_obj = OxmlElement("w:shd")
                    shade_obj.set(qn2("w:fill"),"2ABCB5")
                    table_cell_properties.append(shade_obj)

                    widths = (docx.shared.Inches(7), docx.shared.Inches(0.9))
                    heights = (docx.shared.Inches(1.4), docx.shared.Inches(1.4))

                    for row in programeDescriptionTable.rows:
                        for idx, width in enumerate(widths):
                            row.cells[idx].width = width

                    for idx,row in enumerate(programeDescriptionTable.rows):
                            row.height = heights[idx]

                    for row in programeDescriptionTable.rows:
                        for cell in row.cells:
                            cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.size= docx.shared.Pt(18)

                if self.ableCreator:
                    text = self.programeCreatorE.toPlainText()
                    listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                    final_text = []
                    for i in text:
                        if i in listNubmers:
                            final_text.append(convert_numbers.english_to_hindi(i))
                        else:
                            final_text.append(i)

                    programeCreatorTable = doc.add_table(rows=1,cols=2)

                    programeCreatorTable.style = "Table Grid"
                    hdr_Cells = programeCreatorTable.rows[0].cells

                    hdr_Cells[0].text = "".join(final_text)
                    hdr_Cells[1].text = "المنفذ"
                    programeCreatorTable.autofit = False


                    cell_xml_element = programeCreatorTable.rows[0].cells[1]._tc
                    table_cell_properties = cell_xml_element.get_or_add_tcPr()
                    shade_obj = OxmlElement("w:shd")
                    shade_obj.set(qn2("w:fill"),"2ABCB5")
                    table_cell_properties.append(shade_obj)

                    widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))
                    for row in programeCreatorTable.rows:
                        for idx, width in enumerate(widths):
                            row.cells[idx].width = width

                    for row in programeCreatorTable.rows:
                        for cell in row.cells:
                            cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.size= docx.shared.Pt(18)

            
                if self.ableDate:
                    text = self.programeWhenDateE.toPlainText()
                    listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                    final_text = []
                    for i in text:
                        if i in listNubmers:
                            final_text.append(convert_numbers.english_to_hindi(i))
                        else:
                            final_text.append(i)

                    programeWhenDateTable = doc.add_table(rows=1,cols=2)

                    programeWhenDateTable.style = "Table Grid"
                    hdr_Cells = programeWhenDateTable.rows[0].cells
                    hdr_Cells[1].text = "تاريخ التنفيذ"
                    hdr_Cells[0].text = "".join(final_text)
                    programeWhenDateTable.autofit = False

                    cell_xml_element = programeWhenDateTable.rows[0].cells[1]._tc
                    table_cell_properties = cell_xml_element.get_or_add_tcPr()
                    shade_obj = OxmlElement("w:shd")
                    shade_obj.set(qn2("w:fill"),"2ABCB5")
                    table_cell_properties.append(shade_obj)

                    widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))
                    for row in programeWhenDateTable.rows:
                        for idx, width in enumerate(widths):
                            row.cells[idx].width = width
                    
                    for row in programeWhenDateTable.rows:
                        for cell in row.cells:
                            cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.size= docx.shared.Pt(17)
                if self.ableBenefits:
                    text = self.programeBenefitsE.toPlainText()
                    listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                    final_text = []
                    for i in text:
                        if i in listNubmers:
                            final_text.append(convert_numbers.english_to_hindi(i))
                        else:
                            final_text.append(i)
                            
                    programeBenefitsTable = doc.add_table(rows=1,cols=2)
                    programeBenefitsTable.style = 'Table Grid' #single lines in all cells
                    hdr_Cells = programeBenefitsTable.rows[0].cells
                    hdr_Cells[1].text = "المستفيدون"
                    hdr_Cells[0].text = "".join(final_text)
                    programeBenefitsTable.autofit = False


                    cell_xml_element = programeBenefitsTable.rows[0].cells[1]._tc
                    table_cell_properties = cell_xml_element.get_or_add_tcPr()
                    shade_obj = OxmlElement("w:shd")
                    shade_obj.set(qn2("w:fill"),"2ABCB5")
                    table_cell_properties.append(shade_obj)

                    widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))
                    for row in programeBenefitsTable.rows:
                        for idx, width in enumerate(widths):
                            row.cells[idx].width = width            
                    for row in programeBenefitsTable.rows:
                        for cell in row.cells:
                            cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.size= docx.shared.Pt(15)
                if self.ableCount:
                    text = self.programeNameE.toPlainText()
                    listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                    final_text = []
                    for i in text:
                        if i in listNubmers:
                            final_text.append(convert_numbers.english_to_hindi(i))
                        else:
                            final_text.append(i)

                    programeCountBenefitsTable = doc.add_table(rows=1,cols=2)
                    programeCountBenefitsTable.style = 'Table Grid' 
                    hdr_Cells = programeCountBenefitsTable.rows[0].cells
                    hdr_Cells[1].text = "عدد المستفيدين"
                    hdr_Cells[0].text = "".join(final_text)
                    programeCountBenefitsTable.autofit = False

                    cell_xml_element = programeCountBenefitsTable.rows[0].cells[1]._tc
                    table_cell_properties = cell_xml_element.get_or_add_tcPr()
                    shade_obj = OxmlElement("w:shd")
                    shade_obj.set(qn2("w:fill"),"2ABCB5")
                    table_cell_properties.append(shade_obj)

                    widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))
                    for row in programeCountBenefitsTable.rows:
                        for idx, width in enumerate(widths):
                            row.cells[idx].width = width            
                    for row in programeCountBenefitsTable.rows:
                        for cell in row.cells:
                            cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                            paragraphs = cell.paragraphs
                            for paragraph in paragraphs:
                                for run in paragraph.runs:
                                    font = run.font
                                    font.size= docx.shared.Pt(14)

                if self.countPic != 0:
                    paragraph = doc.add_paragraph()
                    paragraph.paragraph_format.space_before = docx.shared.Pt(2)
                    run = paragraph.add_run()
                    for i in range(len(self.pictersPaths)):
                        if self.pictersPaths[i] !="":
                                try:
                                    os.remove("imageWithBoarder.png")
                                except:
                                    pass
                                imgB = Image.open(self.pictersPaths[i])
                                # this border may cause problems in future
                                border_color_rgb = (128, 128, 128, 255)

                                resize = imgB.resize((500,500),Image.LANCZOS)

                                bordered_image = ImageOps.expand(resize, border=8, fill=border_color_rgb)
                                
                                bordered_image.save('imageWithBoarder.png')
                                if len(self.label1Maye.text()) > 0 or len(self.label2Maye.text()) > 0:
                                    run.add_picture('imageWithBoarder.png',width=docx.shared.Inches(3.9),height=docx.shared.Inches(1.8))
                                else:
                                    run.add_picture('imageWithBoarder.png',width=docx.shared.Inches(3.9),height=docx.shared.Inches(2.2))

                                if i !=1:
                                    run.add_text("   ")
                                if i==1:
                                    run.add_text("\n")
                        paragraph.paragraph_format.space_after = docx.shared.Pt(0)

                if len(self.label1Maye.text()) > 0 or len(self.label2Maye.text()) > 0:
                    addmins_table = doc.add_table(rows=1, cols=2)
                    for row in addmins_table.rows:
                        for cell in row.cells:
                            tc = cell._element.tcPr
                            tc.left = None
                            tc.top = None
                            tc.right = None
                            tc.bottom = None
                            cell.vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            cell.paragraphs[0].alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
                            cell.paragraphs[0].size = docx.shared.Pt(8)

                    addmins_Cells = addmins_table.rows[0].cells

                    addmins_Cells[0].text = self.label1Maye.text()+"\n"+f"{self.consultName.text()}"

                    addmins_table.rows[0].cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

                    addmins_Cells[1].text = self.label2Maye.text()+"\n"+f"{self.MangerName.text()}"

                    addmins_table.rows[0].cells[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

                    paragraph12 =addmins_Cells[0].paragraphs[0]
                    run = paragraph12.runs
                    font = run[0].font
                    font.size= docx.shared.Pt(15)

                    paragraph13 =addmins_Cells[1].paragraphs[0]
                    run = paragraph13.runs
                    font = run[0].font
                    font.size= docx.shared.Pt(15)
                    heights = (docx.shared.Pt(16),docx.shared.Pt(16))
                    for idx,row in enumerate(addmins_table.rows):
                        row.height = heights[idx]
                subFilesD = [f for f in os.listdir(folderFinle) if f.endswith(".docx")]
                name = nameFile+".docx"
                if name in subFilesD:
                    i = 1
                    while name in subFilesD:
                        name = f"({i}) {name}"
                        i+=1

                doc.save(f"{folderFinle}/{name}")
                d = QMessageBox(parent=self.windowCreating,text=f"تم التصدير بنجاح")
                d.setWindowTitle("نجاح")
                d.setIcon(QMessageBox.Icon.Information)
                ret = d.exec()
            try:
                os.remove("pic1")
                os.remove("pic2")
                os.remove("pic4")
                os.remove("pic4")
                os.remove("secretThing.png")
            except:
                pass        



        # Send Pdf to printer after write it
        def printDoc(self):
            try:
                os.remove("printFile.pdf")
            except:
                pass
  
            try:
                self.writePdf(fromWhere="printer")
                webbrowser.open("printFile.pdf", new=2)
                time.sleep(3)
                pyautogui.hotkey("ctrl","p")
            except:
                pass


        # Function To Get Summary _word_ Of Exist Reports
        def exportAllReportsAsWord(self):
                
                self.pdfFilesPaths = []
                cr.execute("SELECT id FROM reports")
                reports = cr.fetchall()
                if reports:
                    self.eachValue = 100//len(reports)
                    self.progressBarWindow = Choices()
                    self.progressBarWindow.setFixedSize(250,30)
                    self.progressBar = QProgressBar(self.progressBarWindow)
                    self.progressBar.setGeometry(0,0,290,30)
                    # choose desktop as dafault path
                    folder_path = os.path.join(os.path.expanduser("~"), "Desktop")
                    # folder_path = QFileDialog.getExistingDirectory(self.windowSaved,"اختر مسارا", desktopPath)
                    if(folder_path):
                        for i in reports:
                          for j in i:
                            self.completeexportAllReportsAsWord(j,folder_path)

                        d = QMessageBox(parent=self.windowCreating,text="تم تصدير الملفات على سطح المكتب بنجاح")
                        d.setWindowTitle("نجاح")
                        d.setIcon(QMessageBox.Icon.Information)
                        d.exec()
                        # the next code is under revision
                        os.execv(sys.executable, ['python'] + sys.argv)
                        app.closeAllWindows()
                    else:
                      pass
                    
                else:
                  d = QMessageBox()
                  d.setText("ليس هناك ملفات للتصدير") 
                  d.setIcon(QMessageBox.Icon.Critical)
                  d.setStandardButtons(QMessageBox.StandardButton.Ok) 
                  d.exec() 


        # Function To Get Summary _pdf_ Of Exist Reports
        def exportAllReportsAsPdf(self):
                
                self.pdfFilesPaths = []
                cr.execute("SELECT id FROM reports")
                reports = cr.fetchall()
                if reports:
                    self.eachValue = 100//len(reports)
                    self.progressBarWindow = Choices()
                    self.progressBarWindow.setFixedSize(250,30)
                    self.progressBar = QProgressBar(self.progressBarWindow)
                    self.progressBar.setGeometry(0,0,290,30)
                    # choose desktop as dafault path
                    folder_path = os.path.join(os.path.expanduser("~"), "Desktop")
                    if(folder_path):
                        for i in reports:
                          for j in i:
                            self.completeexportAllReportsAsPdf(j,folder_path)

                        d = QMessageBox(parent=self.windowCreating,text="تم تصدير الملفات على سطح المكتب بنجاح")
                        d.setWindowTitle("نجاح")
                        d.setIcon(QMessageBox.Icon.Information)
                        d.exec()
                        # the next code is under revision
                        os.execv(sys.executable, ['python'] + sys.argv)
                        app.closeAllWindows()
                    else:
                      pass
                    
                else:
                  d = QMessageBox()
                  d.setText("ليس هناك ملفات للتصدير") 
                  d.setIcon(QMessageBox.Icon.Critical)
                  d.setStandardButtons(QMessageBox.StandardButton.Ok) 
                  d.exec() 


        def completeexportAllReportsAsWord(self,idFun,folder_path):

            try:
                os.remove("pic1.png")
                os.remove("pic2.png")
                os.remove("pic3.png")
                os.remove("pic4.png")
                os.remove("secretThing.png")
            except:
                pass

            doc = docx.Document()
            sections = doc.sections
            for section in sections:
                section.top_margin = docx.shared.Cm(0.7)
                section.bottom_margin = docx.shared.Cm(0.7)
                section.left_margin = docx.shared.Cm(0.7)
                section.right_margin = docx.shared.Cm(0.7)
            sec_pr = doc.sections[0]._sectPr
            pg_borders = OxmlElement('w:pgBorders')
            # specifies how the relative positioning of the borders should be calculated
            pg_borders.set(qn('w:offsetFrom'), 'page')
            for border_name in ('top', 'left', 'bottom', 'right',): # set all borders
                border_el = OxmlElement(f'w:{border_name}')
                border_el.set(qn('w:val'), 'triple') # a single line
                border_el.set(qn('w:sz'), '4') # for meaning of  remaining attrs please look docs
                border_el.set(qn('w:space'), '10')
                border_el.set(qn('w:color'), 'black')
                pg_borders.append(border_el) # register single border to border el
            sec_pr.append(pg_borders) # apply border changes to section

            headers_table = doc.add_table(rows=1, cols=2)
            for row in headers_table.rows:
                for cell in row.cells:
                    tc = cell._element.tcPr
                    tc.left = None
                    tc.top = None
                    tc.right = None
                    tc.bottom = None
                    cell.vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    cell.paragraphs[0].alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
                    cell.paragraphs[0].size = docx.shared.Pt(15)
            hdr_Cells = headers_table.rows[0].cells

            cr.execute("SELECT line1 FROM start")
            hdr_Cells[1].text = cr.fetchone()[0]
            cr.execute("SELECT line2 FROM start")
            hdr_Cells[1].text = hdr_Cells[1].text+"\n"+cr.fetchone()[0]+"\t"
            cr.execute("SELECT line3 FROM start")
            hdr_Cells[1].text = hdr_Cells[1].text+"\n"+cr.fetchone()[0]
            cr.execute("SELECT line4 FROM start")
            hdr_Cells[1].text = hdr_Cells[1].text+"\n"+cr.fetchone()[0]

            widths = (docx.shared.Inches(5.8),docx.shared.Inches(3))
            for row in headers_table.rows:
                for idx, width in enumerate(widths):
                    row.cells[idx].width = width

            heights = (docx.shared.Inches(1.1),docx.shared.Inches(1.1))
            for idx,row in enumerate(headers_table.rows):
                row.height = heights[idx]

            
            paragraph12322 =hdr_Cells[1].paragraphs[0]
            run = paragraph12322.runs
            font = run[0].font
            font.size= docx.shared.Pt(15)

            cells = headers_table.rows[0].cells[0].paragraphs[0]
            runCells = cells.add_run()
            cr.execute(f"SELECT picLogo FROM reports WHERE id={idFun}")
            picLogo = cr.fetchone()[0]
            if picLogo !="":
                with open("secretThing.png","wb") as secretThing:
                    secretThing.write(picLogo)
                runCells.add_picture("secretThing.png",width=docx.shared.Inches(2.1),height=docx.shared.Inches(1))
                runCells.add_text("\t")

            runCells.add_picture("images/logo.png",width=docx.shared.Inches(2.5),height=docx.shared.Inches(1))

            for row in headers_table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
            headers_table.rows[0].cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT


            GoodPrograme = doc.add_paragraph("\t\t\t\t\t\t\tتوثيق برنامج")
            GoodPrograme.runs[0].font.size = docx.shared.Pt(20)
            GoodPrograme.paragraph_format.space_after = docx.shared.Pt(0.1)
            GoodPrograme.paragraph_format.space_before = docx.shared.Pt(1)


            cr.execute(f"SELECT name FROM reports WHERE id={idFun}")  
            namePrograme = cr.fetchone()[0]
            if namePrograme!="":
                text = namePrograme
                listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                final_text = []
                for i in text:
                    if i in listNubmers:
                        final_text.append(convert_numbers.english_to_hindi(i))
                    else:
                        final_text.append(i)

                programeNameProgrameTable = doc.add_table(rows=1,cols=2)
                programeNameProgrameTable.style = "Table Grid"
                hdr_Cells = programeNameProgrameTable.rows[0].cells
                hdr_Cells[1].text = "اسم البرنامج"
                hdr_Cells[0].text = ''.join(final_text)
                programeNameProgrameTable.autofit = False

                cell_xml_element = programeNameProgrameTable.rows[0].cells[1]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement("w:shd")
                shade_obj.set(qn2("w:fill"),"2ABCB5")
                table_cell_properties.append(shade_obj)

                widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))

                for row in programeNameProgrameTable.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width
                
                for row in programeNameProgrameTable.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(17)
                heights = (docx.shared.Inches(.35), docx.shared.Inches(.35))
                for idx,row in enumerate(programeNameProgrameTable.rows):
                        row.height = heights[idx]
            
            cr.execute(f"SELECT Goals FROM reports WHERE id={idFun}")  
            Goals = cr.fetchone()[0]
            if Goals!="":
                text = Goals
                listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                final_text = []
                for i in text:
                    if i in listNubmers:
                        final_text.append(convert_numbers.english_to_hindi(i))
                    else:
                        final_text.append(i)

                programeGolasTable = doc.add_table(rows=1,cols=2)
                programeGolasTable.style = "Table Grid"
                hdr_Cells = programeGolasTable.rows[0].cells
                hdr_Cells[1].text = "\t\tالأهداف"
                hdr_Cells[0].text = ''.join(final_text)
                programeGolasTable.autofit = False

                cell_xml_element = programeGolasTable.rows[0].cells[1]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement("w:shd")
                shade_obj.set(qn2("w:fill"),"2ABCB5")
                table_cell_properties.append(shade_obj)

                widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))
                heights = (docx.shared.Inches(1.4), docx.shared.Inches(1.4))
                for row in programeGolasTable.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width
                for idx,row in enumerate(programeGolasTable.rows):
                    row.height = heights[idx]
                

                for row in programeGolasTable.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(18)
            
            cr.execute(f"SELECT description FROM reports WHERE id={idFun}")  
            description = cr.fetchone()[0]
            if description!="":
                text = description
                listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                final_text = []
                for i in text:
                    if i in listNubmers:
                        final_text.append(convert_numbers.english_to_hindi(i))
                    else:
                        final_text.append(i)

                programeDescriptionTable = doc.add_table(rows=1,cols=2)
                programeDescriptionTable.style = "Table Grid"
                hdr_Cells = programeDescriptionTable.rows[0].cells
                hdr_Cells[1].text = "\t\tالوصف"
                hdr_Cells[0].text = ''.join(final_text)
                programeDescriptionTable.autofit = False

                cell_xml_element = programeDescriptionTable.rows[0].cells[1]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement("w:shd")
                shade_obj.set(qn2("w:fill"),"2ABCB5")
                table_cell_properties.append(shade_obj)

                widths = (docx.shared.Inches(7), docx.shared.Inches(0.9))
                heights = (docx.shared.Inches(1.4), docx.shared.Inches(1.4))

                for row in programeDescriptionTable.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width

                for idx,row in enumerate(programeDescriptionTable.rows):
                        row.height = heights[idx]

                for row in programeDescriptionTable.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(18)
            
            cr.execute(f"SELECT executer FROM reports WHERE id={idFun}")  
            executer = cr.fetchone()[0]
            if executer!="":
                text = executer
                listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                final_text = []
                for i in text:
                    if i in listNubmers:
                        final_text.append(convert_numbers.english_to_hindi(i))
                    else:
                        final_text.append(i)

                programeCreatorTable = doc.add_table(rows=1,cols=2)

                programeCreatorTable.style = "Table Grid"
                hdr_Cells = programeCreatorTable.rows[0].cells

                hdr_Cells[0].text = ''.join(final_text)
                hdr_Cells[1].text = "المنفذ"
                programeCreatorTable.autofit = False


                cell_xml_element = programeCreatorTable.rows[0].cells[1]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement("w:shd")
                shade_obj.set(qn2("w:fill"),"2ABCB5")
                table_cell_properties.append(shade_obj)

                widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))
                for row in programeCreatorTable.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width

                for row in programeCreatorTable.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(18)

            cr.execute(f"SELECT executeDate FROM reports WHERE id={idFun}")
            executeDate = cr.fetchone()[0]
            if executeDate!="":
                text = executeDate
                listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                final_text = []
                for i in text:
                    if i in listNubmers:
                        final_text.append(convert_numbers.english_to_hindi(i))
                    else:
                        final_text.append(i)

                programeWhenDateTable = doc.add_table(rows=1,cols=2)

                programeWhenDateTable.style = "Table Grid"
                hdr_Cells = programeWhenDateTable.rows[0].cells
                hdr_Cells[1].text = "تاريخ التنفيذ"
                hdr_Cells[0].text = ''.join(final_text)
                programeWhenDateTable.autofit = False

                cell_xml_element = programeWhenDateTable.rows[0].cells[1]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement("w:shd")
                shade_obj.set(qn2("w:fill"),"2ABCB5")
                table_cell_properties.append(shade_obj)

                widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))
                for row in programeWhenDateTable.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width
                
                for row in programeWhenDateTable.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(17)
            
            cr.execute(f"SELECT benefits FROM reports WHERE id={idFun}")
            benefits = cr.fetchone()[0]

            if benefits!="":
                text = benefits
                listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                final_text = []
                for i in text:
                    if i in listNubmers:
                        final_text.append(convert_numbers.english_to_hindi(i))
                    else:
                        final_text.append(i)
                
                programeBenefitsTable = doc.add_table(rows=1,cols=2)
                programeBenefitsTable.style = 'Table Grid' 
                hdr_Cells = programeBenefitsTable.rows[0].cells
                hdr_Cells[1].text = "المستفيدون"
                hdr_Cells[0].text = ''.join(final_text)
                programeBenefitsTable.autofit = False


                cell_xml_element = programeBenefitsTable.rows[0].cells[1]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement("w:shd")
                shade_obj.set(qn2("w:fill"),"2ABCB5")
                table_cell_properties.append(shade_obj)

                widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))
                for row in programeBenefitsTable.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width            
                for row in programeBenefitsTable.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(15)
            
            cr.execute(f"SELECT countBenefits FROM reports WHERE id={idFun}")
            countBenefits = cr.fetchone()[0]
            
            if countBenefits!="":
                text = countBenefits
                listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                final_text = []
                for i in text:
                    if i in listNubmers:
                        final_text.append(convert_numbers.english_to_hindi(i))
                    else:
                        final_text.append(i)


                programeCountBenefitsTable = doc.add_table(rows=1,cols=2)
                programeCountBenefitsTable.style = 'Table Grid' 
                hdr_Cells = programeCountBenefitsTable.rows[0].cells
                hdr_Cells[1].text = "عدد المستفيدين"
                hdr_Cells[0].text = ''.join(final_text)
                programeCountBenefitsTable.autofit = False

                cell_xml_element = programeCountBenefitsTable.rows[0].cells[1]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement("w:shd")
                shade_obj.set(qn2("w:fill"),"2ABCB5")
                table_cell_properties.append(shade_obj)

                widths = (docx.shared.Inches(7), docx.shared.Inches(1.1))
                for row in programeCountBenefitsTable.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width            
                for row in programeCountBenefitsTable.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(14)


            self.pictersPaths = []

            cr.execute(f"SELECT pic1 FROM reports WHERE id={idFun}")
            pic1B = cr.fetchone()[0]
            if  pic1B is not None and isinstance(pic1B, bytes):
                with open("pic1.png","wb") as pic1:
                    pic1.write(pic1B)
                self.pictersPaths.append("pic1.png")

            cr.execute(f"SELECT pic2 FROM reports WHERE id={idFun}")
            pic2B = cr.fetchone()[0]
            if  pic2B is not None and isinstance(pic2B, bytes):
                with open("pic2.png","wb") as pic2:
                    pic2.write(pic2B)
                self.pictersPaths.append("pic2.png")

            cr.execute(f"SELECT pic3 FROM reports WHERE id={idFun}")
            pic3B = cr.fetchone()[0]
            if  pic3B is not None and isinstance(pic3B, bytes):
                with open("pic3.png","wb") as pic3:
                    pic3.write(pic3B)
                self.pictersPaths.append("pic3.png")

            cr.execute(f"SELECT pic4 FROM reports WHERE id={idFun}")
            pic4B = cr.fetchone()[0]
            
            if  pic4B is not None and isinstance(pic4B, bytes):
                with open("pic4.png","wb") as pic4:
                    pic4.write(pic4B)
                self.pictersPaths.append("pic4.png")

            if len(self.pictersPaths) > 0:
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.space_before = docx.shared.Pt(2)
                run = paragraph.add_run()
                for i in range(len(self.pictersPaths)):
                    try:
                        os.remove("imageWithBoarder.png")
                    except:
                        pass
                    imgB = Image.open(self.pictersPaths[i])
                    # if you want to add border
                    # border_color_rgb = (128, 128, 128, 255)

                    resize = imgB.resize((500,500),Image.LANCZOS)

                    # bordered_image = ImageOps.expand(resize, border=8, fill=border_color_rgb)
                    bordered_image = ImageOps.expand(resize, border=8)
                    
                    bordered_image.save('imageWithBoarder.png')

                    cr.execute(f"SELECT label1Maybe FROM reports WHERE id={idFun}")
                    label1Maybe = cr.fetchone()[0]

                    cr.execute(f"SELECT label2Maybe FROM reports WHERE id={idFun}")
                    label2Maybe = cr.fetchone()[0]

                    if len(label1Maybe) > 0 or len(label2Maybe) > 0:
                        run.add_picture('imageWithBoarder.png',width=docx.shared.Inches(3.9),height=docx.shared.Inches(1.8))
                    else:
                        run.add_picture('imageWithBoarder.png',width=docx.shared.Inches(3.9),height=docx.shared.Inches(2.2))

                    if i !=1:
                        run.add_text("   ")
                    if i==1:
                        run.add_text("\n")
                    paragraph.paragraph_format.space_after = docx.shared.Pt(0)

            label1Maybe = ""  
            label2Maybe = ""  

            if len(label1Maybe) > 0 or len(label2Maybe) > 0:
                addmins_table = doc.add_table(rows=1, cols=2)
                for row in addmins_table.rows:
                    for cell in row.cells:
                        tc = cell._element.tcPr
                        tc.left = None
                        tc.top = None
                        tc.right = None
                        tc.bottom = None
                        cell.vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        cell.paragraphs[0].alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
                        cell.paragraphs[0].size = docx.shared.Pt(8)

                addmins_Cells = addmins_table.rows[0].cells

                cr.execute(f"SELECT manger FROM reports WHERE id={idFun}")
                manger = cr.fetchone()[0]


                addmins_Cells[0].text = label1Maybe+"\n"+f"{manger}"
                addmins_table.rows[0].cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

                cr.execute(f"SELECT co_manger FROM reports WHERE id={idFun}")
                co_manger = cr.fetchone()[0]

                addmins_Cells[1].text = label2Maybe+"\n"+f"{co_manger}"

                addmins_table.rows[0].cells[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

                paragraph12 =addmins_Cells[0].paragraphs[0]
                run = paragraph12.runs
                font = run[0].font
                font.size= docx.shared.Pt(15)

                paragraph13 =addmins_Cells[1].paragraphs[0]
                run = paragraph13.runs
                font = run[0].font
                font.size= docx.shared.Pt(15)
                
                heights = (docx.shared.Pt(16),docx.shared.Pt(16))
                for idx,row in enumerate(addmins_table.rows):
                    row.height = heights[idx]

            subFilesD = [f for f in os.listdir(folder_path) if f.endswith(".docx")]
            name = str(idFun)+".docx"
            if name in subFilesD:
                i = 1
                while name in subFilesD:
                    name = f"({i}) {name}"
                    i+=1
    
            doc.save(f"{folder_path}/{name}")
            
        
        def completeexportAllReportsAsPdf(self,idFun,folder_path):

            try:
                os.remove("pic1.png")
                os.remove("pic2.png")
                os.remove("pic3.png")
                os.remove("pic4.png")
                os.remove("secretThing.png")
            except:
                pass

            content = [] 
            name = str(idFun)+".pdf"
            pdf_file_path = f"{folder_path}/{name}"
            doc = SimpleDocTemplate(pdf_file_path, pagesize=letter, rightMargin=0, leftMargin=0, topMargin=30, bottomMargin=5)

            # Register the Amiri font
            font_path = 'font/Amiri-Regular.ttf' 
            pdfmetrics.registerFont(TTFont('ArabicFont', font_path))
            font_path_bold = 'font/Amiri-Bold.ttf'  
            pdfmetrics.registerFont(TTFont('ArabicFont-Bold', font_path_bold))
            # Set up styles
            styles = getSampleStyleSheet()
            custom_style = ParagraphStyle('CustomStyle', parent=styles['Normal'], fontSize=16, spaceAfter=14, alignment=1)
            custom_style.fontName = 'ArabicFont-bold'  

            custom_style_header = ParagraphStyle('CustomStyle', parent=styles['Normal'], fontSize=14, alignment=TA_RIGHT)
            custom_style_header.fontName = 'ArabicFont'  

            # Prepare header data with text and placeholders for images
            cr.execute("SELECT line1 FROM start")
            line1=cr.fetchone()[0]
            cr.execute("SELECT line2 FROM start")
            line2=cr.fetchone()[0]
            cr.execute("SELECT line3 FROM start")
            line3=cr.fetchone()[0]
            cr.execute("SELECT line4 FROM start")
            line4=cr.fetchone()[0]
            header_data = [[
                "", 
                "",  
                Paragraph(f"""
                        {
                        get_display(f"{arabic_reshaper.reshape(line1)}")  
                        }
                        {'&nbsp;' * 4}
                        <br/><br/>
                        {
                        get_display(arabic_reshaper.reshape(line2))
                        }
                        {'&nbsp;' * 8}
                        <br/><br/>
                        {
                        get_display(arabic_reshaper.reshape(line3))
                        }
                        <br/><br/>
                        {
                        get_display(arabic_reshaper.reshape(line4))
                        }
                        """, 
                        custom_style_header
                        ), 
                ]]
            
            # Add images to the header
            if os.path.exists("images/logo.png"):
                piclogo =img("images/logo.png", width=160, height=80) 
                header_data[0][1] = piclogo  # Assign the first image to the second column
            
            try:
                cr.execute(f"SELECT picLogo FROM reports WHERE id={idFun}")
                picLogo = cr.fetchone()[0]
                if picLogo !="":
                    with open("secretThing.png","wb") as secretThing:
                       secretThing.write(picLogo)
                    logo = img("secretThing.png", width=165, height=80)  # Adjust dimensions as needed
                    header_data[0][0] = logo  # Assign the second image to the third column
            except:
                pass


            # Calculate dynamic column widths
            page_width = letter[0]    # The width of the letter page
            image_width = 2.6 * inch  # Width allocated for each image column
            max_text_width = page_width - (2.3 * image_width) - 1 * inch  # Remaining width for the first column
            # Create the header table
            header_table = Table(header_data, colWidths=[max_text_width,image_width,image_width])
            header_table.setStyle(TableStyle([
                 ('SIZE', (0, 0), (-1, -1), 14), 
                 ('VALIGN', (1, 0), (1, 0), 'RIGHT'),
                 ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),  # Align all content to the right
                 ('RIGHTPADDING', (0, 0), (-1, -1), -25),  # Remove right padding for all cells
                 ('FONTNAME', (0, 0), (-1, 0),'ArabicFont'),
                 ('BOTTOMPADDING', (0, 0), (-1, 0), 16)              
            ]))

            content.append(header_table)
            # Add additional text (Title)
            title = Paragraph(get_display(arabic_reshaper.reshape("\t\t\t\t\t\t\tتوثيق برنامج")), custom_style)
            content.append(title)

            # *Get Report Content*
            # Style for the section names (right column)
            section_right_style = ParagraphStyle(
                name="rightContent",
                alignment=TA_RIGHT,
                textColor=colors.black,
                fontSize=12,
                spaceAfter=10,
                leading=18
            )
                
            section_right_style.fontName = 'ArabicFont-bold'  # Set the custom style font to ArabicFont
            # Style for the section content (left column)
            section_left_style = ParagraphStyle(
                name="leftContent",
                alignment=TA_RIGHT,
                leading=15,  # Adjust this value to set the desired line spacing
                fontSize=12,
                spaceAfter=10
            )
            section_left_style.fontName = 'ArabicFont' 

            data=[]

            cr.execute(f"SELECT name FROM reports WHERE id={idFun}")  
            namePrograme = cr.fetchone()[0]            
            if namePrograme !="":
                data.append(
                   ((get_display(arabic_reshaper.reshape("اسم البرنامج")), get_display(arabic_reshaper.reshape(namePrograme))))
                )

            cr.execute(f"SELECT Goals FROM reports WHERE id={idFun}")  
            Goals = cr.fetchone()[0]
            if Goals !="":
                goals_text = get_display(arabic_reshaper.reshape(Goals))
                # Ensure line breaks are preserved
                goals_text = goals_text.replace('\n', '<br/>')  # Convert newline characters to <br/>
                data.append(   
                    (get_display(arabic_reshaper.reshape("الأهداف")),goals_text)
                )
            cr.execute(f"SELECT description FROM reports WHERE id={idFun}")  
            description = cr.fetchone()[0]
            if description!="":
                description_text = get_display(arabic_reshaper.reshape(description))
                # Ensure line breaks are preserved
                description_text = description_text.replace('\n', '<br/>')  # Convert newline characters to <br/>
                data.append(   
                    (get_display(arabic_reshaper.reshape("الوصف")),description_text)
                )
         
            cr.execute(f"SELECT executer FROM reports WHERE id={idFun}")  
            executer = cr.fetchone()[0]
            if executer !="":
                data.append(   
                   (get_display(arabic_reshaper.reshape("المنفذ")),get_display(arabic_reshaper.reshape(executer)))
                ) 
            

            cr.execute(f"SELECT executeDate FROM reports WHERE id={idFun}")
            executeDate = cr.fetchone()[0]
            if executeDate !="":
                data.append(   
                   (get_display(arabic_reshaper.reshape("تاريخ التنفيذ")),get_display(arabic_reshaper.reshape(executeDate)))
                   )

            cr.execute(f"SELECT benefits FROM reports WHERE id={idFun}")
            benefits = cr.fetchone()[0]
            if benefits:
                data.append(   
                   (get_display(arabic_reshaper.reshape("المستفيدون")),get_display(arabic_reshaper.reshape(benefits)))
                   ) 
            
            
            cr.execute(f"SELECT countBenefits FROM reports WHERE id={idFun}")
            countBenefits = cr.fetchone()[0]

            if countBenefits !="":
                data.append(   
                   (get_display(arabic_reshaper.reshape("عدد المستفيدين")),get_display(arabic_reshaper.reshape(countBenefits))))
                
        

             # Create the table data by formatting each section name and content
            table_items = []
            for section_name, section_content in data:
                # Right column: Section name (red text)
                right_col = Paragraph(section_name,section_right_style)
                # Left column: Section content
                left_col = Paragraph(section_content,section_left_style)
                # Append the two-column row to table_data
                table_items.append([left_col, right_col])
            table = Table(table_items, colWidths=[6.6 * inch, 1.3 * inch])

            # Add some basic styling to the table (optional)
            table.setStyle(TableStyle([
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),  # Center vertically
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),    # Center horizontally
                ('BACKGROUND', (1, 0), (1, -1), colors.HexColor("#2ABCB5")),  # Right cell background
                ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.black),
                ('BOX', (0, 0), (-1, -1), 0.1, colors.black),
                ('FONTNAME', (0, 0), (-1, 0), 'ArabicFont'),  # Use the registered font name
                ('FONTNAME', (0, 1), (-1, -1), 'ArabicFont'),  # Use the registered font name for data
                ('TOPPADDING', (0, 0), (-1, -1), 3) # Add top padding
            ]))

            content.append(table)

            self.pictersPaths = []

            cr.execute(f"SELECT pic1 FROM reports WHERE id={idFun}")
            pic1B = cr.fetchone()[0]
            if  pic1B is not None and isinstance(pic1B, bytes):
                with open("pic1.png","wb") as pic1:
                    pic1.write(pic1B)
                self.pictersPaths.append("pic1.png")

            cr.execute(f"SELECT pic2 FROM reports WHERE id={idFun}")
            pic2B = cr.fetchone()[0]
            if  pic2B is not None and isinstance(pic2B, bytes):
                with open("pic2.png","wb") as pic2:
                    pic2.write(pic2B)
                self.pictersPaths.append("pic2.png")

            cr.execute(f"SELECT pic3 FROM reports WHERE id={idFun}")
            pic3B = cr.fetchone()[0]
            if  pic3B is not None and isinstance(pic3B, bytes):
                with open("pic3.png","wb") as pic3:
                    pic3.write(pic3B)
                self.pictersPaths.append("pic3.png")

            cr.execute(f"SELECT pic4 FROM reports WHERE id={idFun}")
            pic4B = cr.fetchone()[0]
            
            if  pic4B is not None and isinstance(pic4B, bytes):
                with open("pic4.png","wb") as pic4:
                    pic4.write(pic4B)
                self.pictersPaths.append("pic4.png")

            if len(self.pictersPaths) > 0:

                images_data = []
                max_images_per_row = 2 
                current_row = []

                for i in range(len(self.pictersPaths)):
                    if self.pictersPaths[i] != "":
                        # Calculate dynamic width for images
                        page_width, _ = letter  # Get the width of the page
                        image_width = (page_width / max_images_per_row) * 0.8  # Set image width to 80% of the divided space
                        image = img(self.pictersPaths[i], width=281, height=135)  # Adjust dimensions as needed

                        current_row.append(image)

                        if len(current_row) == max_images_per_row:
                          images_data.append(current_row)  # Add the current row without empty cells
                          current_row = []  # Reset for the next row

                if current_row:
                  images_data.append(current_row)  # Add any remaining images in the last row


                if images_data:
                  # Calculate the total number of columns
                  images_table = Table(images_data, colWidths=[4 * inch, 4 * inch])
                  images_table.setStyle(TableStyle([
                   ('SIZE', (0, 0), (-1, -1), 30),
                   ('FONTNAME', (0, 0), (-1, -1), 'ArabicFont'),
                   ('TOPPADDING', (0, 0), (-1, -1), 2.5),  
                   ('BOTTOMPADDING', (0, 0), (-1, -1), 0),  
                   ('LEFTPADDING', (0, 0), (-1, -1), 3.9  ) 
                   ]))

                  content.append(images_table)
            # *Get Footer Content*
            # Style for  (right column)
            footer_right_style = ParagraphStyle(
                name="rightContent",
                alignment=TA_RIGHT,
                textColor=colors.black,
                fontSize=12,
                spaceAfter=10
            )
            footer_right_style.fontName = 'ArabicFont-bold'  # Set the custom style font to ArabicFont
             # Style for  (left column)
            footer_left_style = ParagraphStyle(
                name="leftContent",
                alignment=TA_LEFT,
                fontSize=12,
                spaceAfter=10
            )
            footer_left_style.fontName = 'ArabicFont-bold' 
            footer_data = [[[],[]],[[],[]]]

            try:
                cr.execute(f"SELECT label2Maybe FROM reports WHERE id={idFun}")
                label2Maybe = cr.fetchone()[0]

                if label2Maybe !="":
                    footer_data[0][0].append(get_display(arabic_reshaper.reshape(self.label2Maye.text())))

                
                cr.execute(f"SELECT co_manger FROM reports WHERE id={idFun}")
                co_manger = cr.fetchone()[0]

                if co_manger !="":
                    footer_data[1][1].append(get_display(arabic_reshaper.reshape(self.consultName.text())))       


                cr.execute(f"SELECT label1Maybe FROM reports WHERE id={idFun}")
                label1Maybe = cr.fetchone()[0]

                if label1Maybe !="":
                    footer_data[0][1].append(get_display(arabic_reshaper.reshape(self.label1Maye.text())))


                cr.execute(f"SELECT manger FROM reports WHERE id={idFun}")
                manger = cr.fetchone()[0]

                if manger !="":
                    footer_data[1][0].append(get_display(arabic_reshaper.reshape(self.MangerName.text())))

                footer_table_items = []
                for first_list, second_list in footer_data:
                    # Convert lists to strings before creating Paragraphs
                    first = ' '.join(first_list) if first_list else ''
                    second = ' '.join(second_list) if second_list else ''
                    right_col = Paragraph(first, footer_right_style)
                    left_col = Paragraph(second, footer_left_style)
                    # Append the two-column row to table_data
                    footer_table_items.append([left_col, right_col])

                # Create the Table
                table = Table(footer_table_items, colWidths=[3.7 * inch, 3.7 * inch])

                # Add some basic styling to the table (optional)
                table.setStyle(TableStyle([
                  ('FONTNAME', (0, 0), (-1, 0), 'ArabicFont'),  # Use the registered font name
                  ('FONTNAME', (0, 1), (-1, -1), 'ArabicFont'),  # Use the registered font name for data
                  ('LEFTPADDING', (0, 0), (-1, -1), 10),  # Add left padding
                   ('RIGHTPADDING', (0, 0), (-1, -1), 10),  # Add right padding
                   ('TOPPADDING', (0, 0), (-1, -1), 10),  # Add top padding
                  ('BOTTOMPADDING', (0, 0), (-1, -1), 10),  # Add bottom padding
                ]))

                content.append(table)
            except:
                  pass


            def add_border(canvas, doc):
                # Draw a border around the page
                width, height = letter  # Get the page dimensions
                border_offset = 20  # Thickness of the border
                canvas.setStrokeColor(colors.black)
                canvas.setLineWidth(1)  # Border width
                canvas.rect(border_offset, border_offset, width - 2 * border_offset, height - 2 * border_offset, stroke=1, fill=0)

            # Build the PDF document
            doc.build(content, onFirstPage=add_border, onLaterPages=add_border) 
              
            

        

        def exportSummaryAsPdf(self):
                try:
                                
                        FileNameSave = QFileDialog.getSaveFileName(self.windowCreating, "اختر مسارا", desktopPath)
                        if len(FileNameSave[0]) > 0:
                            folder = (str(FileNameSave[0]).split("/"))
                            nameFile = folder[-1]
                            folderFinle = "/".join(folder[:-1])
                            cr.execute("SELECT name, executer, executeDate, benefits, countBenefits FROM reports")
                            fetched_data = cr.fetchall()
          
                            header = [
                            get_display(arabic_reshaper.reshape(' م ')),  # Number column header
                            get_display(arabic_reshaper.reshape(' اسم البرنامج ')),  # Program name header
                            get_display(arabic_reshaper.reshape(' المنفذ ')),  # Executer header
                            get_display(arabic_reshaper.reshape(' تاريخ التنفيذ ')),  # Execution date header
                            get_display(arabic_reshaper.reshape(' المستفيدون ')),  # Beneficiaries header
                            get_display(arabic_reshaper.reshape(' عدد المستفيدين '))  # Count of beneficiaries header
                            ]
                        
                            header.reverse()  # Reverse the header for RTL direction
                            data = [header]  # Add header as the first row of the table
                        
                            for numberTemp, row in enumerate(fetched_data):
                                reshaped_row = [                                
                                str(numberTemp + 1),  # Number in Arabic format
                                row[0],  # name
                                row[1],  # Executer
                                row[2],  # Execution Date
                                row[3],  # Beneficiaries
                                row[4],  # Count of beneficiaries
                                 ]
                                reshaped_row.reverse()
                                pdfmetrics.registerFont(TTFont('ArabicFont', 'font/Amiri-Regular.ttf'))
                                styles = getSampleStyleSheet()
                                custom_style = ParagraphStyle(
                                'CustomStyle',
                                parent=styles['Normal'],
                                fontSize=14, 
                                alignment=1     
                                 )
                                custom_style.fontName = 'ArabicFont'
                                
                                # Wrap each cell's text in a Paragraph
                                wrapped_row = [Paragraph(get_display(arabic_reshaper.reshape(text)),custom_style) for text in reshaped_row]
                                data.append(wrapped_row)

                            pdf_file_path = f"{folderFinle}/{nameFile}.pdf"
                            doc = SimpleDocTemplate(pdf_file_path, pagesize=landscape(A4))

                            column_widths = [1.2 * inch, 2.5 * inch, 2.1 * inch, 2.7 * inch, 2.7 * inch,0.3 * inch]  # Adjust as needed
                            # column_widths = [80,120,100,255,240,25]
                            styles = getSampleStyleSheet()
                            custom_style = ParagraphStyle(
                              'CustomStyle',
                               parent=styles['Normal'],
                               fontSize=16,
                               spaceAfter=14,
                               alignment=1 
                            )
                            font_path_bold = 'font/Amiri-Bold.ttf'  
                            pdfmetrics.registerFont(TTFont('ArabicFont-Bold', font_path_bold))
                            # Set the font for the custom style
                            custom_style.fontName = 'ArabicFont-bold'
                            # Create a paragraph before the table
                            paragraph_text = get_display(arabic_reshaper.reshape("ملخص تقارير البرامج"))
                            paragraph = Paragraph(paragraph_text, custom_style)  # Use your desired style
                            
                            table = Table(data, colWidths=column_widths)
                            table.setStyle(TableStyle([
                                
                            #  ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                            #  ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),  # Set vertical alignment to top
                            #  ('FONTNAME', (0, 0), (-1, 0), 'ArabicFont'),  # Use the registered font name
                            #  ('FONTNAME', (0, 1), (-1, -1), 'ArabicFont'),  # Use the registered font name for data
                            #  ('SIZE', (0, 0), (-1, -1), 14),  # Font size
                            #  ('GRID', (0, 0), (-1, -1), 1, colors.black),
                            #  ('LEFTPADDING', (0, 0), (-1, -1), 5),  # Add left padding
                            #  ('RIGHTPADDING', (0, 0), (-1, -1), 5),  # Add right padding
                            #  ('TOPPADDING', (0, 0), (-1, -1), 10),  # Add top padding
                            #  ('BOTTOMPADDING', (0, 0), (-1, -1), 20),  # Add bottom padding
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),  # Center align all
                ('ALIGN', (0, 1), (-1, -1), 'CENTER'),  # Center align headers
                ('ALIGN', (-1, 1), (-1, -1), 'RIGHT'),  # Left align last column (count of beneficiaries)
                ('FONTNAME', (0, 0), (-1, 0), 'ArabicFont'),  # Font for header
                ('FONTNAME', (0, 1), (-1, -1), 'ArabicFont'),  # Font for data
                ('SIZE', (0, 0), (-1, -1), 14),  # Font size
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('LEFTPADDING', (0, 0), (-1, -1), 5),
                ('RIGHTPADDING', (0, 0), (-1, -1), 5),
                ('TOPPADDING', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 20),
                             ]))


                            elements = [paragraph,table]
                            doc.build(elements)

                            d = QMessageBox(parent=self.windowCreating, text=f"تم التصدير بنجاح")
                            d.setWindowTitle("نجاح")
                            d.setIcon(QMessageBox.Icon.Information)
                            d.exec()
   
                except Exception as e:
                    print(e)
                    d = QMessageBox(parent=self.windowCreating)  
                    d.setWindowTitle("فشل")  
                    d.setText("حدث خطأ حاول مرة أخرى")
                    d.setIcon(QMessageBox.Icon.Warning)
                    d.exec() 

        
        def exportSummaryAsWord(self):
            
            FileNameSave = QFileDialog.getSaveFileName(self.windowCreating,"اختر مسارا",desktopPath)
            if len(FileNameSave[0])>0:
                folder = (str(FileNameSave[0]).split(f"/"))
                nameFile = folder[-1]
                folderFinle = "/".join(folder[:-1])
                doc = docx.Document()                
                sections = doc.sections
                for section in sections:
                    section.orientation = docx.enum.section.WD_ORIENTATION.LANDSCAPE
                    section.page_width, section.page_height = section.page_height, section.page_width
                    section.top_margin = docx.shared.Cm(0.7)
                    section.bottom_margin = docx.shared.Cm(0.7)
                    section.left_margin = docx.shared.Cm(0.7)
                    section.right_margin = docx.shared.Cm(0.7)

                # Add a centered bold paragraph before the table
                paragraph = doc.add_paragraph()
                run = paragraph.add_run("ملخص تقارير البرامج")
                run.bold = True  # Make the text bold
                paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # Center the paragraph


                SummryTable = doc.add_table(rows=1,cols=6)
                SummryTable.style = "Table Grid"
                hdr_Cells = SummryTable.rows[0].cells

                hdr_Cells[5].text = "م"
                hdr_Cells[4].text = "اسم البرنامج"
                hdr_Cells[3].text = "المنفذ"
                hdr_Cells[2].text = "تاريخ التنفيذ"
                hdr_Cells[1].text = "المتسفيدون"
                hdr_Cells[0].text = "عدد المستفيدين"

                widths = (docx.shared.Inches(1.5),docx.shared.Inches(1.5),docx.shared.Inches(1.5),docx.shared.Inches(4),docx.shared.Inches(4),docx.shared.Inches(.4))

                cr.execute("SELECT name,executer,executeDate,benefits,countBenefits FROM reports")
                
                for numberTemp,i in enumerate(cr.fetchall()):
                    row_Cells = SummryTable.add_row().cells
                    row_Cells[5].text = str(numberTemp+1)
                    number = 4
                    for j in i:
                        text = j
                        listNubmers = ["1","2","3","4","5","6","7","8","9","0"]
                        final_text = []
                        for xs in text:
                            if xs in listNubmers:
                                final_text.append(convert_numbers.english_to_hindi(xs))
                            else:
                                final_text.append(xs)
                    
                        row_Cells[number].text = "".join(final_text)
                        number-=1
                for row in SummryTable.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = width
                
                for row in SummryTable.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                font = run.font
                                font.size= docx.shared.Pt(17)

                name = nameFile+".docx"
                doc.save(f"{folderFinle}/{name}")
                d = QMessageBox(parent=self.windowCreating,text=f"تم التصدير بنجاح")
                d.setWindowTitle("نجاح")
                d.setIcon(QMessageBox.Icon.Information)
                d.exec()
                
                                
        def closeEvent(self, event):
            try:
                self.sender().objectName()
                event.accept()
            except:
                reply = QMessageBox(self.windowCreating)
                reply.setWindowTitle("تأكيد حفظ")
                reply.setText("هل تريد حفظ التقرير")

                reply.setStandardButtons(QMessageBox.StandardButton.Yes|QMessageBox.StandardButton.No|QMessageBox.StandardButton.Cancel)
                bottonOk = reply.button(QMessageBox.StandardButton.Yes)
                bottonOk.setText("نعم")
    
                bottonCancel = reply.button(QMessageBox.StandardButton.No)
                bottonCancel.setText("لا")
                
                bottonNo = reply.button(QMessageBox.StandardButton.Cancel)
                bottonNo.setText("تم الحفظ")
                x = reply.exec()

                if x == QMessageBox.StandardButton.No or x == QMessageBox.StandardButton.Cancel:
                    event.accept()
                elif x == QMessageBox.StandardButton.Yes:
                    event.ignore()
                    self.SavePrograme()
                    
        def load_data(self):
            # Clear current items
            self.listWidget.clear()
            # Execute the query to fetch reports
            cr.execute("SELECT id, reportName FROM reports")
            for i in cr.fetchall():
                # Create a widget to hold the icon and label
                item_widget = QWidget()
                item_layout = QHBoxLayout()
                # Set layout direction to RsightToLeft
                item_layout.setDirection(QHBoxLayout.Direction.RightToLeft)
                # Create a QPushButton for the icon
                icon_button = QPushButton()
                icon_button.setStyleSheet("Qproperty-icon:url(images/popUpwindow.png); qproperty-iconSize:30px 30px; background-color:transparent")
                icon_button.setFixedSize(40, 40)  # Size of the icon button (adjust as needed)
                icon_button.clicked.connect(lambda event, name=str(i[0]): self.showReport(name))  # Connect to icon click event
                # Create a QLabel for the report name
                label = QLabel(str(i[1]))
                label.setStyleSheet("background-color:transparent; padding: 5px;")
                label.setAlignment(Qt.AlignmentFlag.AlignCenter)
                # Add icon button and label to the layout (icon first in right-to-left)
                item_layout.addWidget(icon_button)
                item_layout.addWidget(label)
                item_layout.addStretch()  # Add a stretch to space them apart

                # Set the layout for the widget
                item_widget.setLayout(item_layout)
                # Create a QListWidgetItem and add the widget to the QListWidget
                list_item = QListWidgetItem(self.listWidget)
                list_item.setSizeHint(item_widget.sizeHint())  # Adjust size based on widget
                list_item.setData(Qt.ItemDataRole.UserRole, i[0])  # Store report ID
                self.listWidget.addItem(list_item)
                self.listWidget.setItemWidget(list_item, item_widget)
                  
    if __name__ == "__main__":
        app = QApplication(sys.argv)
        app.setStyleSheet(
            '''
            QLineEdit{
                font-size:15px
            }
            QLabel{
                font-size:15px
            }
            '''
        )
        window = ReportEditor()
        app.exec()
        
else:
    ctypes.windll.shell32.ShellExecuteW(None, "runas", sys.executable, " ".join(sys.argv), None, 1)
